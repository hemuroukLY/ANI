"""
E2E test for issue-030 kb-service infrastructure (P0 RPCs).

Starts local gateway (dev auth mode) + rag-engine + kb-service, all pointing
at the server-deployed Postgres/Milvus/MinIO/NATS. Tests the P0 RPC flow:
  CreateKB → GetDocumentUploadURL → upload file → NotifyDocumentUploaded
  → (wait parse) → ListChunks → keyword_search → Query → DeleteDocument → DeleteKB

Two file types are tested:
  1. Markdown (.md)  — display text, markdown table, image links
  2. Word (.docx)    — display text, table, embedded images uploaded to MinIO
                       (returns OSS links like [图片: ...](ani-kb-docs/...))

LLM Query is tested with the vLLM service at 10.10.20.181:3011.

Usage:
  python tests/e2e/test_issue030_e2e.py
"""
from __future__ import annotations

import io
import json
import os
import subprocess
import sys
import time
import uuid
from datetime import datetime
from pathlib import Path

# ── paths ────────────────────────────────────────────────────────────────────
REPO = Path(__file__).resolve().parents[4]  # repo/
KB_SVC = REPO / "services" / "kb-service"
RAG_ENGINE = REPO / "ai" / "rag-engine"
GATEWAY_EXE = REPO / "bin" / "ani-gateway.exe"
E2E_LOG = KB_SVC / "tests" / "e2e" / "e2e_result.log"

# ── server component endpoints (NodePorts on 10.10.1.66) ─────────────────────
SRV = "10.10.1.66"
PG_URL = f"postgres://ani:ani_dev_password@{SRV}:30945/ani?sslmode=disable"
MILVUS_ADDR = f"{SRV}:31930"
MINIO_ENDPOINT = f"{SRV}:30900"
MINIO_AK = "ani-s05-minio"
MINIO_SK = "F36UCbnRR-bY9Upv8uuammuBwkHFlTYABiXCbtMCmlc"
NATS_URL = f"nats://{SRV}:31062"

# LLM service (from .env — Qwen3-235B-A22B works, Qwen3.6-35B-A3B returns 500)
VLLM_MODEL = "Qwen3-235B-A22B"
VLLM_API_BASE = "http://10.10.20.181:3011/v1"
VLLM_API_KEY = "sk-YOp8k71BXjxBTeZniPPvQlbGgciH0CB9WOWXkmuCzjfIZ5R8"

# Use the existing test01 tenant (FK constraint on knowledge_bases.tenant_id → tenants.id)
TENANT_ID = "00000000-0000-0000-0000-000000000002"
E2E_TAG = f"e2e-{int(time.time())}"


# ── logging: print to both terminal and file ─────────────────────────────────
_log_fh: io.TextIOBase | None = None


def _open_log():
    global _log_fh
    E2E_LOG.parent.mkdir(parents=True, exist_ok=True)
    _log_fh = open(E2E_LOG, "w", encoding="utf-8")


def log(msg: str = "", **kw):
    """Print to terminal AND log file, with timestamp."""
    ts = datetime.now().strftime("%H:%M:%S")
    line = f"[{ts}] {msg}" if msg else ""
    print(line, flush=True, **kw)
    if _log_fh:
        _log_fh.write(line + "\n")
        _log_fh.flush()


def log_json(label: str, data):
    """Pretty-print a dict/list as JSON to terminal + file."""
    log(f"── {label} ──")
    formatted = json.dumps(data, ensure_ascii=False, indent=2, default=str)
    print(formatted, flush=True)
    if _log_fh:
        _log_fh.write(formatted + "\n\n")
        _log_fh.flush()
    log("")


# ── process management ───────────────────────────────────────────────────────
_procs: list[subprocess.Popen] = []


def _start(cmd: list[str], env: dict[str, str], cwd: Path, name: str):
    """Start a subprocess, streaming stdout/stderr to the log file."""
    log(f"Starting {name}: {' '.join(cmd[:4])}... (cwd={cwd})")
    log_file = open(E2E_LOG.parent / f"{name}.stdout.log", "w", encoding="utf-8")
    proc = subprocess.Popen(
        cmd,
        env=env,
        cwd=str(cwd),
        stdout=log_file,
        stderr=subprocess.STDOUT,
        creationflags=subprocess.CREATE_NEW_PROCESS_GROUP if os.name == "nt" else 0,
    )
    _procs.append(proc)
    return proc


def _wait_http(url: str, label: str, timeout=30):
    """Wait for an HTTP endpoint to respond."""
    import urllib.request
    deadline = time.time() + timeout
    while time.time() < deadline:
        try:
            urllib.request.urlopen(url, timeout=2)
            log(f"  {label} ready at {url}")
            return True
        except Exception:
            time.sleep(1)
    log(f"  {label} NOT ready at {url} (timeout)")
    return False


def _kill_all():
    for p in reversed(_procs):
        try:
            p.terminate()
            p.wait(timeout=5)
        except Exception:
            try:
                p.kill()
            except Exception:
                pass


# ── .docx creation with embedded image ──────────────────────────────────────
def _create_test_docx() -> bytes:
    """Create a .docx file with display text, a table, and an embedded image.

    The embedded image will be extracted by the parse_worker and uploaded to
    MinIO, returning an OSS link like [图片: ...](ani-kb-docs/...).
    """
    from docx import Document
    from docx.shared import Inches, Pt
    import struct

    doc = Document()

    # Display text
    doc.add_heading("E2E Word 测试文档", level=1)
    doc.add_paragraph("这是一段中文显示文字，用于测试Word文档的文本分段能力。")
    doc.add_paragraph("知识库检索是RAG系统的核心组件，支持混合检索模式。")

    # Table
    doc.add_heading("表格内容", level=2)
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "检索模式"
    hdr[1].text = "说明"
    hdr[2].text = "适用场景"
    table.rows[1].cells[0].text = "关键词检索"
    table.rows[1].cells[1].text = "基于pg_trgm的模糊匹配"
    table.rows[1].cells[2].text = "精确术语搜索"
    table.rows[2].cells[0].text = "向量检索"
    table.rows[2].cells[1].text = "基于余弦相似度的语义匹配"
    table.rows[2].cells[2].text = "语义相似搜索"
    table.rows[3].cells[0].text = "混合检索"
    table.rows[3].cells[1].text = "关键词+向量融合"
    table.rows[3].cells[2].text = "综合搜索"

    # Embedded image (generate a minimal PNG in-memory)
    doc.add_heading("嵌入图片", level=2)
    doc.add_paragraph("下方是嵌入文档的图片，将被上传到MinIO并返回OSS链接：")

    # Create a simple 100x100 red PNG image
    import zlib

    def _make_png(width: int, height: int, r: int, g: int, b: int) -> bytes:
        """Generate a minimal valid PNG with a solid color."""
        def _chunk(chunk_type: bytes, data: bytes) -> bytes:
            chunk_data = chunk_type + data
            crc = struct.pack(">I", zlib.crc32(chunk_data) & 0xFFFFFFFF)
            return struct.pack(">I", len(data)) + chunk_data + crc

        # PNG signature
        sig = b"\x89PNG\r\n\x1a\n"
        # IHDR
        ihdr_data = struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0)
        # IDAT — raw pixel data (each row: filter byte + RGB pixels)
        raw = b""
        for _row in range(height):
            raw += b"\x00"  # filter: none
            raw += bytes([r, g, b]) * width
        compressed = zlib.compress(raw)
        # IEND
        iend = b""
        return sig + _chunk(b"IHDR", ihdr_data) + _chunk(b"IDAT", compressed) + _chunk(b"IEND", iend)

    png_bytes = _make_png(100, 100, 220, 50, 50)
    doc.add_picture(io.BytesIO(png_bytes), width=Inches(2))
    doc.add_paragraph("上图: 红色测试图片 (将上传到MinIO)")

    # Add another image
    png_bytes2 = _make_png(80, 80, 50, 180, 50)
    doc.add_picture(io.BytesIO(png_bytes2), width=Inches(1.5))
    doc.add_paragraph("上图: 绿色测试图片 (将上传到MinIO)")

    # Technical details
    doc.add_heading("技术细节", level=2)
    doc.add_paragraph("向量维度为1024，使用BAAI/bge-m3模型进行向量化。")
    doc.add_paragraph("Milvus作为向量数据库，支持HNSW索引和余弦距离度量。")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── service startup ─────────────────────────────────────────────────────────
def start_gateway():
    """Start the Go gateway in dev auth mode, pointing at server components."""
    env = os.environ.copy()
    env.update({
        "ANI_AUTH_MODE": "dev",
        "GATEWAY_PORT": "8080",
        "DATABASE_URL": PG_URL,
        "REDIS_URL": f"redis://:ani_dev_password@{SRV}:30453/0",
        "OBJECT_STORE_PROVIDER": "minio",
        "OBJECT_STORE_ENDPOINT": MINIO_ENDPOINT,
        "OBJECT_STORE_PUBLIC_ENDPOINT": MINIO_ENDPOINT,
        "OBJECT_STORE_ACCESS_KEY_ID": MINIO_AK,
        "OBJECT_STORE_SECRET_ACCESS_KEY": MINIO_SK,
        "OBJECT_STORE_SECURE": "false",
        "OBJECT_STORE_BUCKET_PREFIX": "ani-s13-",
        "OBJECT_STORE_REGION": "us-east-1",
        "VECTOR_STORE_PROVIDER": "milvus",
        "VECTOR_STORE_ENDPOINT": MILVUS_ADDR,
        "VECTOR_STORE_COLLECTION_PREFIX": "ani_s13_",
        "NATS_URL": NATS_URL,
        "STORAGE_PROVIDER": "",
        "K8S_CLUSTER_PROVIDER_MODE": "",
        "GPU_INVENTORY_PROVIDER": "",
        "NETWORK_PROVIDER": "",
        "REGISTRY_PROVIDER_MODE": "",
        "INSTANCE_OBSERVABILITY_PROVIDER": "",
    })
    _start([str(GATEWAY_EXE)], env, REPO / "services" / "ani-gateway", "gateway")


def start_rag_engine():
    """Start the rag-engine with LLM + MinIO image uploader enabled."""
    env = os.environ.copy()
    env.update({
        "MILVUS_ADDR": MILVUS_ADDR,
        "DATABASE_URL": PG_URL,
        "NATS_URL": NATS_URL,
        "EMBEDDING_MODEL": "Qwen3-Embedding-0.6B",
        "EMBEDDING_API_BASE": "http://10.10.20.197:8006/v1",
        # MinIO — for image extraction & upload during parsing
        "MINIO_ENDPOINT": MINIO_ENDPOINT,
        "MINIO_ACCESS_KEY": MINIO_AK,
        "MINIO_SECRET_KEY": MINIO_SK,
        "MINIO_SECURE": "false",
        "MINIO_BUCKET": "ani-kb-docs",
        # LLM — for Query generation
        "VLLM_MODEL": VLLM_MODEL,
        "VLLM_API_BASE": VLLM_API_BASE,
        "VLLM_API_KEY": VLLM_API_KEY,
        "VLLM_CONTEXT_WINDOW": "32768",
        # Redis — for QA service chat store
        "REDIS_URL": f"redis://:ani_dev_password@{SRV}:30453/0",
        # Point rag-engine at the local gateway for object download
        "ANI_GATEWAY_URL": "http://localhost:8080",
    })
    py = sys.executable
    _start([py, "main.py"], env, RAG_ENGINE, "rag-engine")


def start_kb_service():
    """Start kb-service pointing at local gateway + local rag-engine + server DB."""
    env = os.environ.copy()
    env.update({
        "DATABASE_URL": PG_URL,
        "NATS_URL": NATS_URL,
        "ANI_GATEWAY_INTERNAL_URL": "http://localhost:8080",
        "RAG_ENGINE_ADDR": "localhost:8001",
        "GRPC_PORT": "50053",
        # Redis best-effort; use server Redis
        "REDIS_URL": f"redis://:ani_dev_password@{SRV}:30453/0",
    })
    py = sys.executable
    _start([py, "main.py"], env, KB_SVC, "kb-service")


# ── gRPC test client ──────────────────────────────────────────────────────────
def _import_kb_pb():
    """Import the kb-service proto stubs (add generated root to sys.path)."""
    gen_root = str(KB_SVC / "app" / "generated")
    if gen_root not in sys.path:
        sys.path.insert(0, gen_root)
    if str(KB_SVC) not in sys.path:
        sys.path.insert(0, str(KB_SVC))
    from app.generated.kb.v1 import kb_service_pb2 as pb
    from app.generated.kb.v1 import kb_service_pb2_grpc as pb_grpc
    from app.generated.common.v1 import common_pb2
    return pb, pb_grpc, common_pb2


def _upload_file(stub, pb, kb_id: str, file_name: str, file_type: str,
                content_bytes: bytes, idempotency_key: str):
    """Get upload URL, PUT file, return (doc_id, storage_path)."""
    import urllib.request

    upload_req = pb.GetDocumentUploadURLRequest(
        tenant_id=TENANT_ID,
        kb_id=kb_id,
        file_name=file_name,
        file_type=file_type,
        file_size_bytes=0,
        idempotency_key=idempotency_key,
    )
    upload_resp = stub.GetDocumentUploadURL(upload_req)
    doc_id = upload_resp.doc_id
    log_json("GetDocumentUploadURLResponse", {
        "doc_id": doc_id,
        "upload_url": upload_resp.upload_url[:80] + "..." if len(upload_resp.upload_url) > 80 else upload_resp.upload_url,
        "storage_path": upload_resp.storage_path,
    })

    put_req = urllib.request.Request(
        upload_resp.upload_url,
        data=content_bytes,
        method="PUT",
        headers={"Content-Type": "application/octet-stream"},
    )
    put_resp = urllib.request.urlopen(put_req, timeout=10)
    log(f"  Upload PUT status: {put_resp.status} ({len(content_bytes)} bytes)")
    return doc_id, upload_resp.storage_path


def _wait_parse_status(doc_id: str, timeout=120) -> str | None:
    """Wait for parse to reach a terminal state."""
    import asyncpg
    import asyncio

    async def _wait():
        conn = await asyncpg.connect(PG_URL)
        done_states = ("parsed", "done", "completed", "ready", "failed", "error")
        for _ in range(timeout // 2):
            row = await conn.fetchrow(
                "SELECT parse_status, error_message FROM kb_documents WHERE id=$1 AND tenant_id=$2",
                uuid.UUID(doc_id), uuid.UUID(TENANT_ID),
            )
            if row and row["parse_status"] in done_states:
                log(f"  parse_status = {row['parse_status']} error={row['error_message'] or ''}")
                await conn.close()
                return row["parse_status"]
            if row:
                log(f"  parse_status = {row['parse_status']} (still waiting...)")
            await asyncio.sleep(2)
        await conn.close()
        return None

    return asyncio.run(_wait())


def _list_chunks(doc_id: str) -> list[dict]:
    """Fetch all chunks for a document."""
    import asyncpg
    import asyncio

    async def _list():
        conn = await asyncpg.connect(PG_URL)
        rows = await conn.fetch(
            "SELECT id::text, content, chunk_type, content_type, page_number, parent_content "
            "FROM kb_chunks WHERE doc_id=$1 AND tenant_id=$2 ORDER BY chunk_type, created_at",
            uuid.UUID(doc_id), uuid.UUID(TENANT_ID),
        )
        await conn.close()
        return [dict(r) for r in rows]

    return asyncio.run(_list())


def run_e2e():
    """Run the full P0 RPC E2E test."""
    import grpc
    import asyncpg
    import asyncio

    pb, pb_grpc, common_pb2 = _import_kb_pb()

    channel = grpc.insecure_channel("localhost:50053")
    stub = pb_grpc.KBServiceStub(channel)

    results = {"pass": 0, "fail": 0, "errors": []}

    def check(name, cond, detail=""):
        status = "PASS" if cond else "FAIL"
        if cond:
            results["pass"] += 1
        else:
            results["fail"] += 1
            results["errors"].append(f"{name}: {detail}")
        log(f"  [{status}] {name}" + (f" — {detail}" if detail and not cond else ""))

    # ── Test 1: CreateKB ─────────────────────────────────────────────────────
    log("=" * 70)
    log("TEST 1: CreateKB")
    log("=" * 70)
    kb_name = f"e2e-test-{E2E_TAG}"
    req = pb.CreateKBRequest(
        tenant_id=TENANT_ID,
        name=kb_name,
        description="E2E test knowledge base for issue-030",
        embedding_model="bge-m3",
        chunk_size=512,
        top_k=5,
        score_threshold=0.3,
        retrieval_mode="hybrid",
    )
    log_json("CreateKBRequest", {
        "tenant_id": TENANT_ID, "name": kb_name,
        "chunk_size": 512, "top_k": 5, "retrieval_mode": "hybrid",
    })
    try:
        kb = stub.CreateKB(req)
        log_json("CreateKBResponse", {
            "id": kb.id, "name": kb.name, "embedding_model": kb.embedding_model,
            "chunk_size": kb.chunk_size, "top_k": kb.top_k,
            "retrieval_mode": kb.retrieval_mode,
        })
        kb_id = kb.id
        check("CreateKB returns non-empty id", bool(kb_id), "id is empty")
        check("CreateKB name matches", kb.name == kb_name, f"got {kb.name}")
    except grpc.RpcError as e:
        log(f"  CreateKB FAILED: {e.code()}: {e.details()}")
        results["fail"] += 1
        results["errors"].append(f"CreateKB: {e.code()}: {e.details()}")
        channel.close()
        return results

    # ── Test 2: GetKB ─────────────────────────────────────────────────────────
    log("=" * 70)
    log("TEST 2: GetKB")
    log("=" * 70)
    try:
        kb_get = stub.GetKB(pb.GetKBRequest(tenant_id=TENANT_ID, kb_id=kb_id))
        log_json("GetKBResponse", {
            "id": kb_get.id, "name": kb_get.name,
            "retrieval_mode": kb_get.retrieval_mode,
        })
        check("GetKB returns correct id", kb_get.id == kb_id)
    except grpc.RpcError as e:
        log(f"  GetKB FAILED: {e.code()}: {e.details()}")
        results["fail"] += 1
        results["errors"].append(f"GetKB: {e.code()}")

    # ═══════════════════════════════════════════════════════════════════════════
    # File Type 1: Markdown (.md) — display text + table + image links
    # ═══════════════════════════════════════════════════════════════════════════

    # ── Test 3-5: Upload markdown file ────────────────────────────────────────
    log("=" * 70)
    log("TEST 3-5: Upload markdown file (display text + table + image links)")
    log("=" * 70)

    md_content = """# E2E Test Document

## 显示文字段落

这是一段中文显示文字，用于测试知识库的文本分段能力。
知识库检索是RAG系统的核心组件，支持混合检索模式。

## 表格内容

| 检索模式 | 说明 | 适用场景 |
|---------|------|---------|
| 关键词检索 | 基于pg_trgm的模糊匹配 | 精确术语搜索 |
| 向量检索 | 基于余弦相似度的语义匹配 | 语义相似搜索 |
| 混合检索 | 关键词+向量融合 | 综合搜索 |

## 图片链接

文档中包含图片示例：

![架构图](https://example.com/architecture.png)

![流程图](https://example.com/flow.png)

## 技术细节

向量维度为1024，使用BAAI/bge-m3模型进行向量化。
Milvus作为向量数据库，支持HNSW索引和余弦距离度量。
"""
    md_bytes = md_content.encode("utf-8")
    log(f"  Markdown file ({len(md_bytes)} bytes): text + table + image links")
    log_json("MD file content (first 300 chars)", md_content[:300])

    try:
        md_doc_id, md_storage_path = _upload_file(
            stub, pb, kb_id, "e2e_test.md", "md", md_bytes, f"md-{E2E_TAG}",
        )
        check("MD: GetDocumentUploadURL + upload succeeds", True)
    except Exception as e:
        log(f"  MD upload FAILED: {e}")
        check("MD: GetDocumentUploadURL + upload succeeds", False, str(e))
        md_doc_id = None

    # Notify + wait parse for markdown
    if md_doc_id:
        try:
            stub.NotifyDocumentUploaded(pb.NotifyDocumentUploadedRequest(
                tenant_id=TENANT_ID, kb_id=kb_id,
                doc_id=md_doc_id, storage_path=md_storage_path,
            ))
            log("  MD NotifyDocumentUploaded OK")
            check("MD: NotifyDocumentUploaded succeeds", True)
        except grpc.RpcError as e:
            log(f"  MD Notify FAILED: {e.code()}: {e.details()}")
            check("MD: NotifyDocumentUploaded succeeds", False, f"{e.code()}")

        log("  Waiting for MD parse...")
        md_status = _wait_parse_status(md_doc_id)
        check("MD: parse completed", md_status in ("parsed", "done", "completed", "ready"), f"status={md_status}")

    # ── Test 6: ListChunks for markdown (verify text + table + image) ─────────
    log("=" * 70)
    log("TEST 6: ListChunks (markdown) — verify text, table, image links")
    log("=" * 70)
    if md_doc_id:
        md_chunks = _list_chunks(md_doc_id)
        log(f"  Found {len(md_chunks)} chunks for MD doc {md_doc_id}")
        check("MD: chunks written to kb_chunks", len(md_chunks) > 0, "no chunks")

        has_text = has_table = has_image_link = False
        for i, c in enumerate(md_chunks):
            content = c.get("content", "")
            ctype = c.get("chunk_type", "")
            log(f"  ── MD chunk[{i}] type={ctype} content_type={c.get('content_type')} ──")
            log(f"    {content[:200]}{'...' if len(content) > 200 else ''}")
            if "显示文字" in content or "检索" in content:
                has_text = True
            if ("|" in content and "---" in content) or "<table>" in content or "<tr>" in content:
                has_table = True
            if "example.com" in content or "![" in content:
                has_image_link = True

        log(f"  MD content verification: text={has_text} table={has_table} image_link={has_image_link}")
        check("MD: chunks contain display text (显示文字)", has_text)
        check("MD: chunks contain table", has_table)
        check("MD: chunks contain image links", has_image_link)

    # ═══════════════════════════════════════════════════════════════════════════
    # File Type 2: Word (.docx) — display text + table + embedded images → MinIO OSS
    # ═══════════════════════════════════════════════════════════════════════════

    # ── Test 7-9: Upload .docx file with embedded images ──────────────────────
    log("=" * 70)
    log("TEST 7-9: Upload .docx file (text + table + embedded images → MinIO OSS)")
    log("=" * 70)

    docx_bytes = _create_test_docx()
    log(f"  DOCX file ({len(docx_bytes)} bytes): text + table + 2 embedded images")
    log_json("DOCX file info", {"size": len(docx_bytes), "file_name": "e2e_test.docx"})

    try:
        docx_doc_id, docx_storage_path = _upload_file(
            stub, pb, kb_id, "e2e_test.docx", "docx", docx_bytes, f"docx-{E2E_TAG}",
        )
        check("DOCX: GetDocumentUploadURL + upload succeeds", True)
    except Exception as e:
        log(f"  DOCX upload FAILED: {e}")
        check("DOCX: GetDocumentUploadURL + upload succeeds", False, str(e))
        docx_doc_id = None

    # Notify + wait parse for docx
    if docx_doc_id:
        try:
            stub.NotifyDocumentUploaded(pb.NotifyDocumentUploadedRequest(
                tenant_id=TENANT_ID, kb_id=kb_id,
                doc_id=docx_doc_id, storage_path=docx_storage_path,
            ))
            log("  DOCX NotifyDocumentUploaded OK")
            check("DOCX: NotifyDocumentUploaded succeeds", True)
        except grpc.RpcError as e:
            log(f"  DOCX Notify FAILED: {e.code()}: {e.details()}")
            check("DOCX: NotifyDocumentUploaded succeeds", False, f"{e.code()}")

        log("  Waiting for DOCX parse (image extraction → MinIO upload)...")
        docx_status = _wait_parse_status(docx_doc_id)
        check("DOCX: parse completed", docx_status in ("parsed", "done", "completed", "ready"), f"status={docx_status}")

    # ── Test 10: ListChunks for docx (verify text + table + MinIO OSS image links) ──
    log("=" * 70)
    log("TEST 10: ListChunks (docx) — verify text, table, MinIO OSS image links")
    log("=" * 70)
    if docx_doc_id:
        docx_chunks = _list_chunks(docx_doc_id)
        log(f"  Found {len(docx_chunks)} chunks for DOCX doc {docx_doc_id}")
        check("DOCX: chunks written to kb_chunks", len(docx_chunks) > 0, "no chunks")

        has_text = has_table = has_oss_image = False
        oss_urls = []
        for i, c in enumerate(docx_chunks):
            content = c.get("content", "")
            ctype = c.get("chunk_type", "")
            ctype_field = c.get("content_type", "")
            log(f"  ── DOCX chunk[{i}] type={ctype} content_type={ctype_field} ──")
            log(f"    {content[:250]}{'...' if len(content) > 250 else ''}")
            if "显示文字" in content or "检索" in content or "测试" in content:
                has_text = True
            if "<table>" in content or "<tr>" in content or "|" in content:
                has_table = True
            # OSS image links look like [图片: ...](ani-kb-docs/...)
            if "[图片:" in content or "ani-kb-docs/" in content:
                has_oss_image = True
                # Extract OSS URL
                import re
                urls = re.findall(r"\[图片:.*?\]\(([^)]+)\)", content)
                oss_urls.extend(urls)

        log(f"  DOCX content verification: text={has_text} table={has_table} oss_image={has_oss_image}")
        if oss_urls:
            log(f"  Found {len(oss_urls)} MinIO OSS image URLs:")
            for u in oss_urls:
                log(f"    {u}")
        check("DOCX: chunks contain display text", has_text)
        check("DOCX: chunks contain table", has_table)
        check("DOCX: chunks contain MinIO OSS image links ([图片:...](ani-kb-docs/...))", has_oss_image,
              "no [图片:...](ani-kb-docs/...) found in chunks")

    # ── Test 11: keyword_search via repository ─────────────────────────────────
    log("=" * 70)
    log("TEST 11: keyword_search (pg_trgm) — search '向量检索'")
    log("=" * 70)
    sys.path.insert(0, str(KB_SVC))
    from app.repositories import chunk as chunk_repo

    async def _keyword_search():
        conn = await asyncpg.connect(PG_URL)
        try:
            rows = await chunk_repo.keyword_search(
                conn, tenant_id=TENANT_ID, kb_id=kb_id,
                query="向量检索", limit=5,
            )
            return rows
        finally:
            await conn.close()

    kw_results = asyncio.run(_keyword_search())
    log(f"  keyword_search returned {len(kw_results)} results")
    for i, r in enumerate(kw_results[:3]):
        log(f"    [{i}] score={r.get('score', 0):.3f} chunk_id={r.get('chunk_id', '')[:12]}... "
            f"content={r.get('content', '')[:80]}...")
    check("keyword_search returns results", len(kw_results) > 0, "no keyword results")
    if kw_results:
        check("keyword_search has normalized score (0-1)",
              0 <= kw_results[0].get("score", -1) <= 1,
              f"score={kw_results[0].get('score')}")

    # ── Test 12: Query (hybrid retrieval + LLM generation) ─────────────────────
    log("=" * 70)
    log("TEST 12: Query — '什么是知识库检索的检索模式?'")
    log("=" * 70)
    query_req = pb.QueryRequest(
        tenant_id=TENANT_ID,
        kb_id=kb_id,
        question="什么是知识库检索的检索模式?",
        idempotency_key=f"query-{E2E_TAG}",
        top_k=3,
        score_threshold=0.0,
        retrieval_mode="hybrid",
    )
    log_json("QueryRequest", {
        "question": query_req.question,
        "top_k": query_req.top_k,
        "retrieval_mode": query_req.retrieval_mode,
    })
    try:
        query_resp = stub.Query(query_req, timeout=120)
        log_json("QueryResponse", {
            "answer": query_resp.answer[:500] if query_resp.answer else "",
            "session_id": query_resp.session_id,
            "num_sources": len(query_resp.sources),
            "input_tokens": query_resp.input_tokens,
            "output_tokens": query_resp.output_tokens,
        })
        check("Query returns an answer", bool(query_resp.answer), "answer is empty")
        if query_resp.sources:
            for i, sc in enumerate(query_resp.sources[:3]):
                log(f"    source[{i}] score={sc.score:.3f} content={sc.content[:100]}...")
        check("Query returns source chunks", len(query_resp.sources) > 0, "no sources")
    except grpc.RpcError as e:
        err_detail = e.details() or ""
        log(f"  Query FAILED: {e.code()}: {err_detail}")
        check("Query returns an answer", False, f"{e.code()}: {err_detail[:200]}")

    # ── Test 13: DeleteDocument (markdown) ─────────────────────────────────────
    log("=" * 70)
    log("TEST 13: DeleteDocument (markdown)")
    log("=" * 70)
    if md_doc_id:
        try:
            stub.DeleteDocument(pb.DeleteDocumentRequest(
                tenant_id=TENANT_ID, kb_id=kb_id, doc_id=md_doc_id,
            ))
            log("  MD DeleteDocument OK")
            check("MD: DeleteDocument succeeds", True)
        except grpc.RpcError as e:
            log(f"  MD Delete FAILED: {e.code()}: {e.details()}")
            check("MD: DeleteDocument succeeds", False, f"{e.code()}")

        # Verify chunks deleted
        async def _check_md_chunks():
            conn = await asyncpg.connect(PG_URL)
            count = await conn.fetchval(
                "SELECT count(*) FROM kb_chunks WHERE doc_id=$1 AND tenant_id=$2",
                uuid.UUID(md_doc_id), uuid.UUID(TENANT_ID),
            )
            await conn.close()
            return count
        remaining = asyncio.run(_check_md_chunks())
        check("MD: chunks deleted after DeleteDocument", remaining == 0, f"{remaining} chunks remain")

    # ── Test 14: DeleteDocument (docx) ─────────────────────────────────────────
    log("=" * 70)
    log("TEST 14: DeleteDocument (docx)")
    log("=" * 70)
    if docx_doc_id:
        try:
            stub.DeleteDocument(pb.DeleteDocumentRequest(
                tenant_id=TENANT_ID, kb_id=kb_id, doc_id=docx_doc_id,
            ))
            log("  DOCX DeleteDocument OK")
            check("DOCX: DeleteDocument succeeds", True)
        except grpc.RpcError as e:
            log(f"  DOCX Delete FAILED: {e.code()}: {e.details()}")
            check("DOCX: DeleteDocument succeeds", False, f"{e.code()}")

        # Verify chunks deleted
        async def _check_docx_chunks():
            conn = await asyncpg.connect(PG_URL)
            count = await conn.fetchval(
                "SELECT count(*) FROM kb_chunks WHERE doc_id=$1 AND tenant_id=$2",
                uuid.UUID(docx_doc_id), uuid.UUID(TENANT_ID),
            )
            await conn.close()
            return count
        remaining = asyncio.run(_check_docx_chunks())
        check("DOCX: chunks deleted after DeleteDocument", remaining == 0, f"{remaining} chunks remain")

    # ── Test 15: DeleteKB ──────────────────────────────────────────────────────
    log("=" * 70)
    log("TEST 15: DeleteKB")
    log("=" * 70)
    try:
        stub.DeleteKB(pb.DeleteKBRequest(tenant_id=TENANT_ID, kb_id=kb_id))
        log("  DeleteKB OK")
        check("DeleteKB succeeds", True)
    except grpc.RpcError as e:
        log(f"  DeleteKB FAILED: {e.code()}: {e.details()}")
        check("DeleteKB succeeds", False, f"{e.code()}")

    channel.close()
    return results


# ── main ──────────────────────────────────────────────────────────────────────
def main():
    _open_log()
    log("=" * 70)
    log("  ANI Issue-030 E2E Test")
    log(f"  Tenant: {TENANT_ID}")
    log(f"  Tag:    {E2E_TAG}")
    log(f"  Log:    {E2E_LOG}")
    log("=" * 70)
    log("")
    log("Server components (NodePorts on 10.10.1.66):")
    log(f"  Postgres:  {SRV}:30945")
    log(f"  Milvus:    {SRV}:31930")
    log(f"  MinIO:     {SRV}:30900")
    log(f"  NATS:      {SRV}:31062")
    log(f"  Redis:     {SRV}:30453")
    log(f"  Embedding: 10.10.20.197:8006")
    log(f"  LLM:       {VLLM_API_BASE} ({VLLM_MODEL})")
    log("")
    log("Test file types:")
    log("  1. Markdown (.md)  — display text, table, image links")
    log("  2. Word (.docx)   — display text, table, embedded images → MinIO OSS")
    log("")

    # 1. Start services
    log(">>> Phase 1: Starting local services")
    start_gateway()
    if not _wait_http("http://localhost:8080/healthz", "gateway", 15):
        log("FATAL: gateway not ready, aborting")
        _kill_all()
        return 1

    start_rag_engine()
    if not _wait_http("http://localhost:8001/health", "rag-engine", 60):
        log("WARNING: rag-engine not ready (continuing — gRPC may still start)")

    start_kb_service()
    if not _wait_http("http://localhost:8002/health", "kb-service", 15):
        log("FATAL: kb-service not ready, aborting")
        _kill_all()
        return 1

    log("")
    log(">>> Phase 2: Running P0 RPC tests")
    log("")
    time.sleep(2)  # let gRPC server bind

    try:
        results = run_e2e()
    except Exception as e:
        import traceback
        log(f"E2E test crashed: {e}")
        traceback.print_exc()
        results = {"pass": 0, "fail": 1, "errors": [str(e)]}

    log("")
    log("=" * 70)
    log(f"  E2E RESULTS: {results['pass']} passed, {results['fail']} failed")
    log("=" * 70)
    if results["errors"]:
        log("Errors:")
        for e in results["errors"]:
            log(f"  - {e}")
    log("")
    log(f"Full log: {E2E_LOG}")

    _kill_all()
    return 1 if results["fail"] > 0 else 0


if __name__ == "__main__":
    sys.exit(main())
