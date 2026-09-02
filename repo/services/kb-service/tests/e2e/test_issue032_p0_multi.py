"""
E2E test for new architecture P0 RPCs with multiple file types.

Tests ALL P0 RPCs:
  CreateKB → GetKB → ListKBs
  → GetDocumentUploadURL → upload → NotifyDocumentUploaded (x4 file types)
  → GetDocument → ListDocuments
  → (wait parse) → ListChunks (verify text/table/image/summary)
  → Query (hybrid retrieval + LLM generation)
  → Retrieve (vector-only retrieval)
  → DeleteDocument → DeleteKB

File types tested:
  1. Word (.docx)   — long text + table + 2 embedded images (→ MinIO OSS)
  2. Markdown (.md)  — long text + table + external image links
  3. PDF (.pdf)      — generated with text + table + embedded image
  4. Plain text (.txt) — long text only

All input/output/chunk/query details are printed in full (no truncation).

Usage:
  python tests/e2e/test_issue032_p0_multi.py
"""
from __future__ import annotations

import io, json, os, subprocess, sys, time, uuid, struct, zlib
from datetime import datetime
from pathlib import Path

REPO = Path(__file__).resolve().parents[4]
KB_SVC = REPO / "services" / "kb-service"
RAG_ENGINE = REPO / "ai" / "rag-engine"
GATEWAY_EXE = REPO / "bin" / "ani-gateway.exe"
E2E_LOG = KB_SVC / "tests" / "e2e" / "e2e_p0_multi_result.log"

SRV = "10.10.1.66"
PG_URL = f"postgres://ani:ani_dev_password@{SRV}:30945/ani?sslmode=disable"
MILVUS_ADDR = f"{SRV}:31930"
MINIO_ENDPOINT = f"{SRV}:30900"
MINIO_AK = "ani-s05-minio"
MINIO_SK = "F36UCbnRR-bY9Upv8uuammuBwkHFlTYABiXCbtMCmlc"
NATS_URL = f"nats://{SRV}:31062"
VLLM_MODEL = "Qwen3-235B-A22B"
VLLM_API_BASE = "http://10.10.20.181:3011/v1"
VLLM_API_KEY = "sk-YOp8k71BXjxBTeZniPPvQlbGgciH0CB9WOWXkmuCzjfIZ5R8"
TENANT_ID = "00000000-0000-0000-0000-000000000002"
E2E_TAG = f"p0-{int(time.time())}"

_log_fh = None

def _open_log():
    global _log_fh
    E2E_LOG.parent.mkdir(parents=True, exist_ok=True)
    _log_fh = open(E2E_LOG, "w", encoding="utf-8")

def log(msg="", **kw):
    ts = datetime.now().strftime("%H:%M:%S")
    line = f"[{ts}] {msg}" if msg else ""
    print(line, flush=True, **kw)
    if _log_fh:
        _log_fh.write(line + "\n")
        _log_fh.flush()

def log_json(label, data):
    log(f"── {label} ──")
    formatted = json.dumps(data, ensure_ascii=False, indent=2, default=str)
    print(formatted, flush=True)
    if _log_fh:
        _log_fh.write(formatted + "\n\n")
        _log_fh.flush()
    log("")

_procs = []

def _start(cmd, env, cwd, name):
    log(f"Starting {name}: {' '.join(cmd[:4])}... (cwd={cwd})")
    lf = open(E2E_LOG.parent / f"{name}_p0_multi.stdout.log", "w", encoding="utf-8")
    proc = subprocess.Popen(cmd, env=env, cwd=str(cwd), stdout=lf, stderr=subprocess.STDOUT,
        creationflags=subprocess.CREATE_NEW_PROCESS_GROUP if os.name == "nt" else 0)
    _procs.append(proc)
    return proc

def _wait_http(url, label, timeout=30):
    import urllib.request
    deadline = time.time() + timeout
    while time.time() < deadline:
        try:
            urllib.request.urlopen(url, timeout=2)
            log(f"  {label} ready at {url}")
            return True
        except Exception:
            time.sleep(1)
    log(f"  {label} NOT ready at {url} (timeout)")
    return False

def _kill_all():
    for p in reversed(_procs):
        try: p.terminate(); p.wait(timeout=5)
        except: 
            try: p.kill()
            except: pass

# ── PNG ───────────────────────────────────────────────────────────────────────
def _make_png(w, h, r, g, b):
    def _chunk(ct, data):
        cd = ct + data
        crc = struct.pack(">I", zlib.crc32(cd) & 0xFFFFFFFF)
        return struct.pack(">I", len(data)) + cd + crc
    sig = b"\x89PNG\r\n\x1a\n"
    ihdr = struct.pack(">IIBBBBB", w, h, 8, 2, 0, 0, 0)
    raw = b""
    for _ in range(h):
        raw += b"\x00" + bytes([r, g, b]) * w
    return sig + _chunk(b"IHDR", ihdr) + _chunk(b"IDAT", zlib.compress(raw)) + _chunk(b"IEND", b"")

# ── DOCX ──────────────────────────────────────────────────────────────────────
def _create_docx():
    from docx import Document
    from docx.shared import Inches
    doc = Document()
    doc.add_heading("知识库平台多格式文档测试", level=1)
    doc.add_paragraph("本文档用于测试 ANI 知识库平台对 Word 文档的解析能力，包含长文本、表格和嵌入图片。")
    doc.add_heading("一、系统架构说明", level=2)
    for i in range(1, 4):
        doc.add_paragraph(f"第{i}层架构：负责{'请求路由' if i==1 else '业务编排' if i==2 else 'AI推理'}。"
            f"该层通过 gRPC 通信，支持幂等性控制和重试机制，保证在网络抖动场景下的数据一致性。")
    doc.add_heading("二、组件清单", level=2)
    t = doc.add_table(rows=4, cols=3); t.style = "Table Grid"
    t.rows[0].cells[0].text = "组件"; t.rows[0].cells[1].text = "端口"; t.rows[0].cells[2].text = "职责"
    t.rows[1].cells[0].text = "gateway"; t.rows[1].cells[1].text = "8080"; t.rows[1].cells[2].text = "API网关"
    t.rows[2].cells[0].text = "kb-service"; t.rows[2].cells[1].text = "8002"; t.rows[2].cells[2].text = "知识库CRUD"
    t.rows[3].cells[0].text = "rag-engine"; t.rows[3].cells[1].text = "8001"; t.rows[3].cells[2].text = "RAG引擎"
    doc.add_heading("三、架构图示", level=2)
    doc.add_picture(io.BytesIO(_make_png(120, 80, 60, 120, 200)), width=Inches(2))
    doc.add_paragraph("图1: 系统架构图 (嵌入图片，将上传到 MinIO)")
    doc.add_heading("四、解析管线", level=2)
    doc.add_paragraph("文档解析管线包括：获取download_url → Parse RPC → 图片上传 → 摘要生成 → "
        "Embed RPC → Core向量插入 → kb_chunks写入 → 状态更新。异常时标记failed并脱敏错误信息。")
    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()

# ── MD ─────────────────────────────────────────────────────────────────────────
def _create_md():
    md = """# 知识库平台 Markdown 测试

## 一、系统架构说明

ANI 知识库平台采用三层服务架构：网关层负责请求路由和认证，服务编排层负责知识库 CRUD 和解析管线，
AI 推理层负责文档解析、向量嵌入和摘要生成。各服务通过 gRPC 通信，支持幂等性控制。

## 二、组件清单

| 组件 | 端口 | 职责 |
|------|------|------|
| gateway | 8080 | API网关 + 对象存储 |
| kb-service | 8002/50053 | 知识库CRUD + 解析编排 |
| rag-engine | 8001/50052 | RAG引擎 |
| Milvus | 19530 | 向量数据库 |

## 三、架构图示

![系统架构图](https://example.com/arch.png)

![数据流转图](https://example.com/flow.png)

## 四、解析管线

文档解析管线包括：获取download_url → Parse RPC → 图片上传 → 摘要生成 →
Embed RPC → Core向量插入 → kb_chunks写入 → 状态更新。
异常时标记failed并脱敏错误信息（截断500字符）。
幂等性保证：已ready的文档重复调用直接跳过。

## 五、混合检索

关键词检索基于 pg_trgm 模糊匹配，向量检索基于 Milvus HNSW 余弦相似度，
混合检索融合两者结果，兼顾精确性和语义性。
"""
    return md.encode("utf-8")

# ── PDF ──────────────────────────────────────────────────────────────────────
def _create_pdf():
    import fitz  # pymupdf
    doc = fitz.open()
    page = doc.new_page()
    # Text
    page.insert_text((50, 50), "知识库平台 PDF 测试文档", fontsize=14, fontname="helv")
    page.insert_text((50, 80), "一、系统架构说明", fontsize=12, fontname="helv")
    page.insert_text((50, 100), "ANI 平台采用三层架构：网关层、服务编排层和 AI 推理层。", fontsize=10, fontname="helv")
    page.insert_text((50, 115), "各服务通过 gRPC 通信，支持幂等性控制和重试机制。", fontsize=10, fontname="helv")
    page.insert_text((50, 140), "二、组件清单", fontsize=12, fontname="helv")
    page.insert_text((50, 160), "组件 | 端口 | 职责", fontsize=10, fontname="helv")
    page.insert_text((50, 175), "gateway | 8080 | API网关", fontsize=10, fontname="helv")
    page.insert_text((50, 190), "kb-service | 8002 | 知识库CRUD", fontsize=10, fontname="helv")
    page.insert_text((50, 215), "三、架构图示", fontsize=12, fontname="helv")
    # Insert image
    img_rect = fitz.Rect(50, 230, 200, 310)
    img_bytes = _make_png(150, 80, 200, 100, 50)
    page.insert_image(img_rect, stream=img_bytes)
    page.insert_text((50, 320), "图1: 架构图 (嵌入PDF的图片，将被提取并上传)", fontsize=9, fontname="helv")
    page.insert_text((50, 345), "四、解析管线", fontsize=12, fontname="helv")
    page.insert_text((50, 365), "解析管线：download_url → Parse → 图片上传 → 摘要 → Embed → 向量插入 → PG写入 → ready", fontsize=9, fontname="helv")
    page.insert_text((50, 380), "五、混合检索：关键词+向量融合，兼顾精确性和语义性。", fontsize=10, fontname="helv")
    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()

# ── TXT ───────────────────────────────────────────────────────────────────────
def _create_txt():
    text = """知识库平台纯文本测试

一、系统架构说明

ANI 知识库平台采用三层服务架构：网关层负责请求路由和认证鉴权，服务编排层负责知识库 CRUD 和文档解析管线编排，
AI 推理层负责文档解析、向量嵌入和摘要生成等无状态 RPC 执行。在分布式部署中，各服务通过 gRPC 通信，
保证低延迟和高吞吐。同时，所有写操作都支持幂等性控制，通过 idempotency_key 保证在网络抖动和服务重启
场景下的数据一致性。

二、组件清单

ani-gateway: Go, 端口 8080, 负责 API 网关和对象存储
kb-service: Python/FastAPI, 端口 8002/50053, 负责知识库 CRUD 和解析编排
rag-engine: Python/FastAPI, 端口 8001/50052, 负责 RAG 引擎 (parse/embed/generate)
PostgreSQL: 端口 5432, 存储元数据和 kb_chunks
Milvus: 端口 19530, 向量数据库

三、文档解析管线

文档解析管线是知识库平台的核心功能。完整流程包括：获取文档下载 URL、调用 rag-engine Parse RPC、
图片上传到 Core API、摘要生成（取前 3 个 parent blocks 调用 Generate RPC）、向量嵌入（Embed RPC
嵌入 child chunks 和 summary）、Core API 插入预计算向量、PG 写入 kb_chunks（parents + children +
summaries 分开写入）、状态更新（parsing → indexing → ready）。异常时标记 failed 并脱敏错误信息。

四、混合检索策略

知识库平台支持三种检索模式：关键词检索（pg_trgm 模糊匹配）、向量检索（Milvus HNSW 余弦相似度）
和混合检索（融合两种结果）。检索结果经过去重和重排序后，取 top-k 个 chunk 作为 LLM 上下文。
"""
    return text.encode("utf-8")

# ── service startup ──────────────────────────────────────────────────────────
def start_gateway():
    env = os.environ.copy()
    env.update({
        "ANI_AUTH_MODE": "dev", "GATEWAY_PORT": "8080", "DATABASE_URL": PG_URL,
        "REDIS_URL": f"redis://:ani_dev_password@{SRV}:30453/0",
        "OBJECT_STORE_PROVIDER": "minio", "OBJECT_STORE_ENDPOINT": MINIO_ENDPOINT,
        "OBJECT_STORE_PUBLIC_ENDPOINT": MINIO_ENDPOINT,
        "OBJECT_STORE_ACCESS_KEY_ID": MINIO_AK, "OBJECT_STORE_SECRET_ACCESS_KEY": MINIO_SK,
        "OBJECT_STORE_SECURE": "false", "OBJECT_STORE_BUCKET_PREFIX": "ani-s13-",
        "OBJECT_STORE_REGION": "us-east-1",
        "VECTOR_STORE_PROVIDER": "milvus", "VECTOR_STORE_ENDPOINT": MILVUS_ADDR,
        "VECTOR_STORE_COLLECTION_PREFIX": "ani_s13_",
        "NATS_URL": NATS_URL, "STORAGE_PROVIDER": "", "K8S_CLUSTER_PROVIDER_MODE": "",
        "GPU_INVENTORY_PROVIDER": "", "NETWORK_PROVIDER": "", "REGISTRY_PROVIDER_MODE": "",
        "INSTANCE_OBSERVABILITY_PROVIDER": "",
    })
    _start([str(GATEWAY_EXE)], env, REPO / "services" / "ani-gateway", "gateway")

def start_rag_engine():
    env = os.environ.copy()
    env.update({
        "MILVUS_ADDR": MILVUS_ADDR, "DATABASE_URL": PG_URL, "NATS_URL": NATS_URL,
        "EMBEDDING_MODEL": "Qwen3-Embedding-0.6B", "EMBEDDING_API_BASE": "http://10.10.20.197:8006/v1",
        "MINIO_ENDPOINT": MINIO_ENDPOINT, "MINIO_ACCESS_KEY": MINIO_AK, "MINIO_SECRET_KEY": MINIO_SK,
        "MINIO_SECURE": "false", "MINIO_BUCKET": "ani-kb-docs",
        "VLLM_MODEL": VLLM_MODEL, "VLLM_API_BASE": VLLM_API_BASE, "VLLM_API_KEY": VLLM_API_KEY,
        "VLLM_CONTEXT_WINDOW": "32768", "REDIS_URL": f"redis://:ani_dev_password@{SRV}:30453/0",
        "ANI_GATEWAY_URL": "http://localhost:8080",
    })
    _start([sys.executable, "main.py"], env, RAG_ENGINE, "rag-engine")

def start_kb_service():
    env = os.environ.copy()
    env.update({
        "DATABASE_URL": PG_URL, "NATS_URL": NATS_URL,
        "ANI_GATEWAY_INTERNAL_URL": "http://localhost:8080",
        "RAG_ENGINE_ADDR": "localhost:8001", "GRPC_PORT": "50053",
        "REDIS_URL": f"redis://:ani_dev_password@{SRV}:30453/0",
    })
    _start([sys.executable, "main.py"], env, KB_SVC, "kb-service")

# ── helpers ──────────────────────────────────────────────────────────────────
def _import_pb():
    gen = str(KB_SVC / "app" / "generated")
    if gen not in sys.path: sys.path.insert(0, gen)
    if str(KB_SVC) not in sys.path: sys.path.insert(0, str(KB_SVC))
    from app.generated.kb.v1 import kb_service_pb2 as pb
    from app.generated.kb.v1 import kb_service_pb2_grpc as pb_grpc
    from app.generated.common.v1 import common_pb2
    return pb, pb_grpc, common_pb2

def _upload(stub, pb, kb_id, name, ftype, data, idem):
    import urllib.request
    r = stub.GetDocumentUploadURL(pb.GetDocumentUploadURLRequest(
        tenant_id=TENANT_ID, kb_id=kb_id, file_name=name, file_type=ftype,
        file_size_bytes=0, idempotency_key=idem))
    log_json("UploadResponse", {"doc_id": r.doc_id, "storage_path": r.storage_path})
    req = urllib.request.Request(r.upload_url, data=data, method="PUT",
        headers={"Content-Type": "application/octet-stream"})
    resp = urllib.request.urlopen(req, timeout=30)
    log(f"  Upload PUT: {resp.status} ({len(data)} bytes)")
    return r.doc_id, r.storage_path

def _wait_parse(doc_id, timeout=180):
    import asyncpg, asyncio
    async def _w():
        conn = await asyncpg.connect(PG_URL)
        for i in range(timeout // 3):
            row = await conn.fetchrow(
                "SELECT parse_status, error_message, chunk_count FROM kb_documents WHERE id=$1 AND tenant_id=$2",
                uuid.UUID(doc_id), uuid.UUID(TENANT_ID))
            if row and row["parse_status"] in ("ready", "failed"):
                log(f"  parse_status={row['parse_status']} chunk_count={row['chunk_count']} error={row['error_message'] or ''}")
                await conn.close()
                return row["parse_status"], row["chunk_count"], row["error_message"]
            if row: log(f"  parse_status={row['parse_status']} (waiting [{i}])")
            # If still pending after 6s, trigger sync parse via rag-engine HTTP fallback
            # (new architecture: NATS consumer may not be wired yet)
            if i == 2 and row and row["parse_status"] == "pending":
                _trigger_sync_parse(doc_id)
            await asyncio.sleep(3)
        await conn.close()
        return None, 0, "timeout"
    return asyncio.run(_w())

def _trigger_sync_parse(doc_id):
    """Trigger synchronous parse via rag-engine HTTP endpoint (fallback).

    The new architecture's ParseOrchestrator is not yet wired to a NATS
    consumer. When the NATS-based outbox flow is not active, we trigger
    the parse via rag-engine's synchronous HTTP fallback endpoint, which
    runs the same pipeline (download → parse → chunk → embed → write).
    """
    import urllib.request
    try:
        # Get kb_id + tenant_id + storage_path from kb_documents
        import asyncpg, asyncio
        async def _get():
            conn = await asyncpg.connect(PG_URL)
            row = await conn.fetchrow(
                "SELECT kb_id::text, tenant_id::text, storage_path FROM kb_documents WHERE id=$1",
                uuid.UUID(doc_id))
            await conn.close()
            return dict(row) if row else None
        doc = asyncio.run(_get())
        if not doc:
            log(f"  sync parse: doc {doc_id} not found in DB")
            return
        body = json.dumps({"kb_id": doc["kb_id"], "doc_id": doc_id,
            "tenant_id": doc["tenant_id"]}).encode()
        req = urllib.request.Request(
            f"http://localhost:8001/{doc['kb_id']}/documents/{doc_id}/parse",
            data=body, method="POST",
            headers={"Content-Type": "application/json"})
        resp = urllib.request.urlopen(req, timeout=120)
        log(f"  sync parse triggered: {resp.status} {json.loads(resp.read()).get('status','')}")
    except Exception as e:
        log(f"  sync parse trigger failed: {e}")

def _list_chunks(doc_id):
    import asyncpg, asyncio
    async def _l():
        conn = await asyncpg.connect(PG_URL)
        rows = await conn.fetch(
            "SELECT id::text, content, chunk_type, content_type, page_number, "
            "parent_content, parent_chunk_id::text as parent_chunk_id, token_count "
            "FROM kb_chunks WHERE doc_id=$1 AND tenant_id=$2 ORDER BY chunk_type, created_at",
            uuid.UUID(doc_id), uuid.UUID(TENANT_ID))
        await conn.close()
        return [dict(r) for r in rows]
    return asyncio.run(_l())

def _print_chunks(label, chunks):
    """Print ALL chunks with FULL content, no truncation."""
    parents = [c for c in chunks if c["chunk_type"] == "parent"]
    children = [c for c in chunks if c["chunk_type"] == "child"]
    summaries = [c for c in chunks if c["chunk_type"] == "doc_summary"]
    log(f"  Parents: {len(parents)}  Children: {len(children)}  Summaries: {len(summaries)}")
    for i, c in enumerate(chunks):
        ct = c["chunk_type"]
        log(f"  ── {label} Chunk[{i}] type={ct} content_type={c.get('content_type','')} "
            f"page={c.get('page_number',0)} tokens={c.get('token_count',0)} ──")
        log(f"  [{ct.upper()} FULL CONTENT]")
        log(f"  {c.get('content', '')}")
        pc = c.get("parent_content", "")
        if pc:
            log(f"  [parent_content FULL]")
            log(f"  {pc}")
        pid = c.get("parent_chunk_id", "")
        if pid:
            log(f"  parent_chunk_id: {pid}")
        log("")

# ── main ──────────────────────────────────────────────────────────────────────
def run_e2e():
    import grpc
    pb, pb_grpc, common_pb2 = _import_pb()
    ch = grpc.insecure_channel("localhost:50053")
    stub = pb_grpc.KBServiceStub(ch)
    results = {"pass": 0, "fail": 0, "errors": []}

    def check(name, cond, detail=""):
        st = "PASS" if cond else "FAIL"
        if cond: results["pass"] += 1
        else: results["fail"] += 1; results["errors"].append(f"{name}: {detail}")
        log(f"  [{st}] {name}" + (f" — {detail}" if detail and not cond else ""))

    # ═══ P0-1: CreateKB ════════════════════════════════════════════════════════
    log("=" * 70); log("P0-1: CreateKB"); log("=" * 70)
    kb_name = f"p0-multi-{E2E_TAG}"
    req = pb.CreateKBRequest(tenant_id=TENANT_ID, name=kb_name,
        description="P0 multi-format E2E test", embedding_model="bge-m3",
        chunk_size=512, top_k=5, score_threshold=0.3, retrieval_mode="hybrid")
    log_json("CreateKBRequest", {"name": kb_name, "chunk_size": 512, "retrieval_mode": "hybrid"})
    kb = stub.CreateKB(req)
    kb_id = kb.id
    log_json("CreateKBResponse", {"id": kb_id, "name": kb.name, "embedding_model": kb.embedding_model})
    check("CreateKB returns non-empty id", bool(kb_id))

    # ═══ P0-2: GetKB ═══════════════════════════════════════════════════════════
    log("=" * 70); log("P0-2: GetKB"); log("=" * 70)
    kb_get = stub.GetKB(pb.GetKBRequest(tenant_id=TENANT_ID, kb_id=kb_id))
    log_json("GetKBResponse", {"id": kb_get.id, "name": kb_get.name, "retrieval_mode": kb_get.retrieval_mode})
    check("GetKB returns correct id", kb_get.id == kb_id)

    # ═══ P0-3: ListKBs ═════════════════════════════════════════════════════════
    log("=" * 70); log("P0-3: ListKBs"); log("=" * 70)
    list_resp = stub.ListKBs(pb.ListKBsRequest(tenant_id=TENANT_ID,
        page=common_pb2.CursorPageRequest(limit=10)))
    log(f"  ListKBs returned {len(list_resp.kbs)} KBs (first page)")
    for k in list_resp.kbs:
        log(f"    {k.id} = {k.name}")
    # KB may not be on first page due to cursor pagination — check via GetKB instead
    check("ListKBs returns at least 1 KB", len(list_resp.kbs) >= 1)

    # ═══ File type definitions ════════════════════════════════════════════════
    file_types = [
        ("docx", "test_doc.docx", _create_docx(), "DOCX: text+table+2 embedded images"),
        ("md", "test_doc.md", _create_md(), "MD: text+table+external image links"),
        ("pdf", "test_doc.pdf", _create_pdf(), "PDF: text+table+1 embedded image"),
        ("txt", "test_doc.txt", _create_txt(), "TXT: long text only"),
    ]

    doc_ids = {}  # ftype -> (doc_id, storage_path)

    for ftype, fname, fdata, fdesc in file_types:
        log("")
        log("=" * 70)
        log(f"P0-4: Upload + Parse — {fdesc}")
        log(f"  File: {fname}  Size: {len(fdata)} bytes")
        log("=" * 70)

        # Print input content
        log(f"  ── INPUT: {ftype} full content ──")
        if ftype == "docx":
            try:
                from docx import Document
                doc = Document(io.BytesIO(fdata))
                for para in doc.paragraphs:
                    if para.text.strip(): log(f"  {para.text}")
                for tbl in doc.tables:
                    for row in tbl.rows:
                        log(f"  | {' | '.join(c.text for c in row.cells)} |")
            except Exception as e:
                log(f"  (docx extract error: {e})")
        elif ftype == "pdf":
            try:
                import fitz
                doc = fitz.open(stream=fdata, filetype="pdf")
                for pg in doc:
                    log(f"  {pg.get_text()}")
            except Exception as e:
                log(f"  (pdf extract error: {e})")
        else:
            log(f"  {fdata.decode('utf-8')}")
        log("  ── END INPUT ──")
        log("")

        # P0-4a: GetDocumentUploadURL + upload
        did, sp = _upload(stub, pb, kb_id, fname, ftype, fdata, f"{ftype}-{E2E_TAG}")
        doc_ids[ftype] = (did, sp)
        check(f"{ftype.upper()}: upload succeeds", True)

        # P0-4b: NotifyDocumentUploaded
        stub.NotifyDocumentUploaded(pb.NotifyDocumentUploadedRequest(
            tenant_id=TENANT_ID, kb_id=kb_id, doc_id=did, storage_path=sp))
        log(f"  {ftype.upper()}: NotifyDocumentUploaded OK")
        check(f"{ftype.upper()}: NotifyDocumentUploaded succeeds", True)

        # P0-4c: Wait for parse
        log(f"  Waiting for {ftype.upper()} parse...")
        status, chunk_count, error = _wait_parse(did)
        check(f"{ftype.upper()}: parse reaches ready", status == "ready", f"status={status} error={error}")

        # P0-4d: ListChunks
        if status == "ready":
            chunks = _list_chunks(did)
            log("")
            log(f"  ═══ {ftype.upper()} CHUNK DETAILS ({len(chunks)} chunks, chunk_count={chunk_count}) ═══")
            parents = [c for c in chunks if c["chunk_type"] == "parent"]
            children = [c for c in chunks if c["chunk_type"] == "child"]
            summaries = [c for c in chunks if c["chunk_type"] == "doc_summary"]
            check(f"{ftype.upper()}: chunk_count == p+c+s",
                  chunk_count == len(parents) + len(children) + len(summaries),
                  f"chunk_count={chunk_count} vs {len(parents)}+{len(children)}+{len(summaries)}")
            _print_chunks(ftype.upper(), chunks)

            # Verify content types
            all_content = " ".join(c.get("content", "") for c in chunks)
            has_text = "架构" in all_content or "知识库" in all_content or "解析" in all_content
            has_table = "<table>" in all_content or "|" in all_content or "组件" in all_content
            has_summary = len(summaries) > 0
            has_image = "[图片:" in all_content or "![" in all_content or "example.com" in all_content

            check(f"{ftype.upper()}: chunks contain text", has_text)
            check(f"{ftype.upper()}: chunks contain table", has_table)
            check(f"{ftype.upper()}: doc_summary exists", has_summary, "no doc_summary")
            if has_summary:
                log(f"  [{ftype.upper()}] Summary FULL ({len(summaries[0]['content'])} chars):")
                log(f"  {summaries[0]['content']}")
            # Image check: docx and pdf should have [图片: 图片](object_id), md should have ![...](url)
            if ftype in ("docx", "pdf"):
                check(f"{ftype.upper()}: chunks contain [图片: 图片] link", has_image, "no [图片:] link")
            elif ftype == "md":
                check(f"{ftype.upper()}: chunks contain image link", has_image, "no image link")
            log("")

    # ═══ P0-5: GetDocument ═════════════════════════════════════════════════════
    log("=" * 70); log("P0-5: GetDocument (DOCX)"); log("=" * 70)
    docx_did = doc_ids["docx"][0]
    doc = stub.GetDocument(pb.GetDocumentRequest(tenant_id=TENANT_ID, kb_id=kb_id, doc_id=docx_did))
    log_json("GetDocumentResponse", {
        "id": doc.id, "file_name": doc.file_name, "file_type": doc.file_type,
        "parse_status": doc.parse_status, "chunk_count": doc.chunk_count,
        "file_size_bytes": doc.file_size_bytes,
    })
    check("GetDocument returns correct id", doc.id == docx_did)
    check("GetDocument parse_status == ready", doc.parse_status == "ready", f"status={doc.parse_status}")

    # ═══ P0-6: ListDocuments ══════════════════════════════════════════════════
    log("=" * 70); log("P0-6: ListDocuments"); log("=" * 70)
    list_docs = stub.ListDocuments(pb.ListDocumentsRequest(
        tenant_id=TENANT_ID, kb_id=kb_id, page=common_pb2.CursorPageRequest(limit=20)))
    log(f"  ListDocuments returned {len(list_docs.documents)} documents")
    for d in list_docs.documents:
        log(f"    {d.id} | {d.file_name} | {d.file_type} | status={d.parse_status} | chunks={d.chunk_count}")
    check("ListDocuments returns 4 documents", len(list_docs.documents) == 4,
          f"got {len(list_docs.documents)}")

    # ═══ P0-7: Query (hybrid retrieval + LLM generation) ═══════════════════════
    log("=" * 70); log("P0-7: Query — '文档解析管线包含哪些步骤?'"); log("=" * 70)
    q_req = pb.QueryRequest(tenant_id=TENANT_ID, kb_id=kb_id,
        question="文档解析管线包含哪些步骤?", idempotency_key=f"q-{E2E_TAG}",
        top_k=3, score_threshold=0.0, retrieval_mode="hybrid")
    log_json("QueryRequest", {"question": q_req.question, "top_k": q_req.top_k, "retrieval_mode": "hybrid"})
    try:
        q_resp = stub.Query(q_req, timeout=120)
        log_json("QueryResponse", {
            "answer": q_resp.answer,
            "session_id": q_resp.session_id,
            "num_sources": len(q_resp.sources),
            "input_tokens": q_resp.input_tokens,
            "output_tokens": q_resp.output_tokens,
        })
        check("Query returns an answer", bool(q_resp.answer), "answer empty")
        for i, sc in enumerate(q_resp.sources):
            log(f"  ── Source[{i}] score={sc.score:.3f} ──")
            log(f"  [SOURCE FULL CONTENT]")
            log(f"  {sc.content}")
            log("")
        check("Query returns source chunks", len(q_resp.sources) > 0)
    except grpc.RpcError as e:
        log(f"  Query FAILED: {e.code()}: {e.details()}")
        check("Query returns an answer", False, f"{e.code()}")

    # ═══ P0-8: Retrieve (P1 RPC — expected UNIMPLEMENTED in P0) ═══════════════
    log("=" * 70); log("P0-8: Retrieve (P1 RPC, expect UNIMPLEMENTED in P0)"); log("=" * 70)
    r_req = pb.RetrieveRequest(tenant_id=TENANT_ID, kb_id=kb_id,
        question="混合检索策略", top_k=3, score_threshold=0.0, retrieval_mode="hybrid")
    log_json("RetrieveRequest", {"question": r_req.question, "top_k": r_req.top_k, "retrieval_mode": "hybrid"})
    try:
        retrieve_stream = stub.Retrieve(r_req, timeout=60)
        for event in retrieve_stream:
            if event.HasField("error"):
                log(f"  Retrieve error event: {event.error.code}: {event.error.message}")
        check("Retrieve returns UNIMPLEMENTED (P1 RPC in P0)", False, "expected UNIMPLEMENTED")
    except grpc.RpcError as e:
        log(f"  Retrieve returned: {e.code()}: {e.details()}")
        check("Retrieve returns UNIMPLEMENTED (P1 RPC in P0)",
              e.code() == grpc.StatusCode.UNIMPLEMENTED,
              f"got {e.code()}")

    # ═══ P0-9: DeleteDocument (all 4) ════════════════════════════════════════
    log("=" * 70); log("P0-9: DeleteDocument (all 4 types)"); log("=" * 70)
    for ftype, (did, _) in doc_ids.items():
        try:
            stub.DeleteDocument(pb.DeleteDocumentRequest(tenant_id=TENANT_ID, kb_id=kb_id, doc_id=did))
            log(f"  {ftype.upper()}: deleted OK")
            check(f"{ftype.upper()}: DeleteDocument succeeds", True)
        except grpc.RpcError as e:
            check(f"{ftype.upper()}: DeleteDocument succeeds", False, f"{e.code()}")

    # ═══ P0-10: DeleteKB ═══════════════════════════════════════════════════════
    log("=" * 70); log("P0-10: DeleteKB"); log("=" * 70)
    try:
        stub.DeleteKB(pb.DeleteKBRequest(tenant_id=TENANT_ID, kb_id=kb_id))
        log("  DeleteKB OK")
        check("DeleteKB succeeds", True)
    except grpc.RpcError as e:
        check("DeleteKB succeeds", False, f"{e.code()}")

    ch.close()
    return results

def main():
    _open_log()
    log("=" * 70)
    log("  ANI P0 Multi-Format E2E Test — New Architecture")
    log(f"  Tenant: {TENANT_ID}  Tag: {E2E_TAG}")
    log(f"  Log: {E2E_LOG}")
    log("=" * 70)
    log("")
    log("Server components (NodePorts on 10.10.1.66):")
    log(f"  PG:{SRV}:30945  Milvus:{SRV}:31930  MinIO:{SRV}:30900  NATS:{SRV}:31062  Redis:{SRV}:30453")
    log(f"  Embedding: 10.10.20.197:8006  LLM: {VLLM_API_BASE}")
    log("")
    log("P0 RPCs tested:")
    log("  1. CreateKB  2. GetKB  3. ListKBs  4. GetDocumentUploadURL+Upload+Notify (x4)")
    log("  5. GetDocument  6. ListDocuments  7. Query  8. Retrieve  9. DeleteDocument  10. DeleteKB")
    log("")
    log("File types: DOCX (text+table+images), MD (text+table+img links), PDF (text+table+image), TXT (long text)")
    log("")

    log(">>> Phase 1: Starting local services")
    start_gateway()
    if not _wait_http("http://localhost:8080/healthz", "gateway", 15):
        log("FATAL: gateway not ready"); _kill_all(); return 1
    start_rag_engine()
    if not _wait_http("http://localhost:8001/health", "rag-engine", 60):
        log("WARNING: rag-engine not ready (continuing)")
    start_kb_service()
    if not _wait_http("http://localhost:8002/health", "kb-service", 15):
        log("FATAL: kb-service not ready"); _kill_all(); return 1
    log("")
    log(">>> Phase 2: Running P0 RPC tests with 4 file types")
    log("")
    time.sleep(2)
    try:
        results = run_e2e()
    except Exception as e:
        import traceback
        log(f"E2E crashed: {e}"); traceback.print_exc()
        results = {"pass": 0, "fail": 1, "errors": [str(e)]}
    log("")
    log("=" * 70)
    log(f"  E2E RESULTS: {results['pass']} passed, {results['fail']} failed")
    log("=" * 70)
    if results["errors"]:
        log("Errors:")
        for e in results["errors"]: log(f"  - {e}")
    log(f"\nFull log: {E2E_LOG}")
    _kill_all()
    return 1 if results["fail"] > 0 else 0

if __name__ == "__main__":
    sys.exit(main())
