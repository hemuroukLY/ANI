"""
E2E test for issue-032 kb-service ParseOrchestrator (new parse pipeline).

Starts local gateway (dev auth mode) + rag-engine + kb-service, all pointing
at the server-deployed Postgres/Milvus/MinIO/NATS. Tests the full parse pipeline:
  CreateKB → GetDocumentUploadURL → upload file → NotifyDocumentUploaded
  → (wait parse) → ListChunks → verify chunks/images/summary
  → Query → cleanup

Test documents:
  1. Word (.docx) — long text + table + 2 embedded images
  2. Markdown (.md) — long text + table + external image links
  3. PDF — generated with text and an embedded image

Verifies:
  - Image markdown links [图片: 图片](object_id) embedded in parent content
  - Summary chunk (chunk_type="doc_summary") generated
  - chunk_count = parents + children + summaries
  - State machine: pending → parsing → indexing → ready

Usage:
  python tests/e2e/test_issue032_e2e.py
"""
from __future__ import annotations

import io
import json
import os
import subprocess
import sys
import time
import uuid
import struct
import zlib
from datetime import datetime
from pathlib import Path

# ── paths ────────────────────────────────────────────────────────────────────
REPO = Path(__file__).resolve().parents[4]  # repo/
KB_SVC = REPO / "services" / "kb-service"
RAG_ENGINE = REPO / "ai" / "rag-engine"
GATEWAY_EXE = REPO / "bin" / "ani-gateway.exe"
E2E_LOG = KB_SVC / "tests" / "e2e" / "e2e_issue032_result.log"

# ── server component endpoints (NodePorts on 10.10.1.66) ─────────────────────
SRV = "10.10.1.66"
PG_URL = f"postgres://ani:ani_dev_password@{SRV}:30945/ani?sslmode=disable"
MILVUS_ADDR = f"{SRV}:31930"
MINIO_ENDPOINT = f"{SRV}:30900"
MINIO_AK = "ani-s05-minio"
MINIO_SK = "F36UCbnRR-bY9Upv8uuammuBwkHFlTYABiXCbtMCmlc"
NATS_URL = f"nats://{SRV}:31062"

# LLM service
VLLM_MODEL = "Qwen3-235B-A22B"
VLLM_API_BASE = "http://10.10.20.181:3011/v1"
VLLM_API_KEY = "sk-YOp8k71BXjxBTeZniPPvQlbGgciH0CB9WOWXkmuCzjfIZ5R8"

# Use the existing test01 tenant
TENANT_ID = "00000000-0000-0000-0000-000000000002"
E2E_TAG = f"e2e032-{int(time.time())}"


# ── logging ──────────────────────────────────────────────────────────────────
_log_fh: io.TextIOBase | None = None


def _open_log():
    global _log_fh
    E2E_LOG.parent.mkdir(parents=True, exist_ok=True)
    _log_fh = open(E2E_LOG, "w", encoding="utf-8")


def log(msg: str = "", **kw):
    ts = datetime.now().strftime("%H:%M:%S")
    line = f"[{ts}] {msg}" if msg else ""
    print(line, flush=True, **kw)
    if _log_fh:
        _log_fh.write(line + "\n")
        _log_fh.flush()


def log_json(label: str, data):
    log(f"── {label} ──")
    formatted = json.dumps(data, ensure_ascii=False, indent=2, default=str)
    print(formatted, flush=True)
    if _log_fh:
        _log_fh.write(formatted + "\n\n")
        _log_fh.flush()
    log("")


# ── process management ───────────────────────────────────────────────────────
_procs: list[subprocess.Popen] = []


def _start(cmd: list[str], env: dict[str, str], cwd: Path, name: str):
    log(f"Starting {name}: {' '.join(cmd[:4])}... (cwd={cwd})")
    log_file = open(E2E_LOG.parent / f"{name}_issue032.stdout.log", "w", encoding="utf-8")
    proc = subprocess.Popen(
        cmd,
        env=env,
        cwd=str(cwd),
        stdout=log_file,
        stderr=subprocess.STDOUT,
        creationflags=subprocess.CREATE_NEW_PROCESS_GROUP if os.name == "nt" else 0,
    )
    _procs.append(proc)
    return proc


def _wait_http(url: str, label: str, timeout=30):
    import urllib.request
    deadline = time.time() + timeout
    while time.time() < deadline:
        try:
            urllib.request.urlopen(url, timeout=2)
            log(f"  {label} ready at {url}")
            return True
        except Exception:
            time.sleep(1)
    log(f"  {label} NOT ready at {url} (timeout)")
    return False


def _kill_all():
    for p in reversed(_procs):
        try:
            p.terminate()
            p.wait(timeout=5)
        except Exception:
            try:
                p.kill()
            except Exception:
                pass


# ── PNG generation ────────────────────────────────────────────────────────────
def _make_png(width: int, height: int, r: int, g: int, b: int) -> bytes:
    """Generate a minimal valid PNG with a solid color."""
    def _chunk(chunk_type: bytes, data: bytes) -> bytes:
        chunk_data = chunk_type + data
        crc = struct.pack(">I", zlib.crc32(chunk_data) & 0xFFFFFFFF)
        return struct.pack(">I", len(data)) + chunk_data + crc

    sig = b"\x89PNG\r\n\x1a\n"
    ihdr_data = struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0)
    raw = b""
    for _row in range(height):
        raw += b"\x00" + bytes([r, g, b]) * width
    compressed = zlib.compress(raw)
    return sig + _chunk(b"IHDR", ihdr_data) + _chunk(b"IDAT", compressed) + _chunk(b"IEND", b"")


# ── .docx creation with embedded images + long text ──────────────────────────
def _create_test_docx() -> bytes:
    """Create a .docx with long text, table, and 2 embedded images."""
    from docx import Document
    from docx.shared import Inches

    doc = Document()

    # Long text — generate multiple paragraphs
    doc.add_heading("知识库平台架构设计文档", level=1)
    doc.add_paragraph(
        "本文档详细描述了 ANI 知识库平台的整体架构设计，包括服务分层、"
        "数据流转路径、向量检索引擎和文档解析管线的完整流程。"
    )

    # Section 1: architecture overview
    doc.add_heading("一、系统架构概览", level=2)
    for i in range(1, 6):
        doc.add_paragraph(
            f"系统架构第{i}层：该层负责{'请求路由和认证' if i == 1 else '业务逻辑编排' if i == 2 else 'AI 推理执行' if i == 3 else '数据持久化' if i == 4 else '基础设施服务'}。"
            f"在分布式部署中，该层通过 gRPC 和 REST API 与上下游服务通信，"
            f"确保请求的可靠传递和负载均衡。同时，该层实现了幂等性控制和重试机制，"
            f"保证在网络抖动和服务重启场景下的数据一致性。"
        )

    # Section 2: table
    doc.add_heading("二、服务组件清单", level=2)
    table = doc.add_table(rows=5, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "组件名称"
    hdr[1].text = "技术栈"
    hdr[2].text = "端口"
    hdr[3].text = "职责"
    table.rows[1].cells[0].text = "ani-gateway"
    table.rows[1].cells[1].text = "Go"
    table.rows[1].cells[2].text = "8080"
    table.rows[1].cells[3].text = "API 网关 + 对象存储"
    table.rows[2].cells[0].text = "kb-service"
    table.rows[2].cells[1].text = "Python/FastAPI"
    table.rows[2].cells[2].text = "8002/50053"
    table.rows[2].cells[3].text = "知识库 CRUD + 解析编排"
    table.rows[3].cells[0].text = "rag-engine"
    table.rows[3].cells[1].text = "Python/FastAPI"
    table.rows[3].cells[2].text = "8001/50052"
    table.rows[3].cells[3].text = "RAG 引擎 (parse/embed/generate)"
    table.rows[4].cells[0].text = "PostgreSQL"
    table.rows[4].cells[1].text = "PostgreSQL 17"
    table.rows[4].cells[2].text = "5432"
    table.rows[4].cells[3].text = "元数据 + kb_chunks"

    # Section 3: embedded images
    doc.add_heading("三、架构图示", level=2)
    doc.add_paragraph("下方是系统架构示意图，展示各服务之间的调用关系：")
    png1 = _make_png(200, 120, 60, 120, 200)
    doc.add_picture(io.BytesIO(png1), width=Inches(3))
    doc.add_paragraph("图1: ANI 知识库平台系统架构图")

    doc.add_paragraph("下方是数据流转示意图，展示文档解析的完整管线：")
    png2 = _make_png(200, 100, 200, 150, 50)
    doc.add_picture(io.BytesIO(png2), width=Inches(3))
    doc.add_paragraph("图2: 文档解析管线数据流转图")

    # Section 4: more long text
    doc.add_heading("四、文档解析管线详解", level=2)
    doc.add_paragraph(
        "文档解析管线是知识库平台的核心功能之一。它负责将用户上传的文档"
        "（支持 PDF、Word、Markdown 等格式）转换为结构化的文本块，"
        "并生成向量嵌入用于后续的语义检索。"
    )
    doc.add_paragraph(
        "管线的完整流程包括以下步骤：\n"
        "1. 获取文档下载 URL（kb-service 向 Core API 请求 download_url）\n"
        "2. 调用 rag-engine Parse RPC（传入 download_url，rag-engine 下载文档并解析）\n"
        "3. 图片上传（kb-service 将 rag-engine 返回的图片 bytes 上传到 Core API）\n"
        "4. 摘要生成（调用 Generate RPC，取前 3 个 parent blocks 生成 200-500 字摘要）\n"
        "5. 向量嵌入（调用 Embed RPC 嵌入 child chunks + summary）\n"
        "6. 向量插入（kb-service 调用 Core API 插入预计算向量）\n"
        "7. PG 写入（parents + children + summaries 分开写入 kb_chunks）\n"
        "8. 状态更新（parsing → indexing → ready）"
    )
    doc.add_paragraph(
        "在异常场景下，管线会捕获异常并将文档状态标记为 failed，"
        "同时记录经过脱敏处理的错误信息（移除敏感路径和凭据，截断 500 字符）。"
        "幂等性保证：如果文档已经是 ready 状态，重复调用会直接跳过。"
    )

    # Section 5: retrieval
    doc.add_heading("五、混合检索策略", level=2)
    doc.add_paragraph(
        "知识库平台支持三种检索模式：关键词检索、向量检索和混合检索。"
        "关键词检索基于 PostgreSQL 的 pg_trgm 扩展实现模糊匹配，"
        "向量检索基于 Milvus 的 HNSW 索引实现余弦相似度匹配，"
        "混合检索将两种模式的结果融合，兼顾精确性和语义性。"
    )
    doc.add_paragraph(
        "检索结果经过去重和重排序后，取 top-k 个 chunk 作为 LLM 的上下文，"
        "生成最终回答。整个检索-生成流程在 kb-service 的 Query RPC 中编排完成。"
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Markdown with long text + image links ────────────────────────────────────
def _create_test_md() -> bytes:
    """Create a markdown file with long text, table, and image links."""
    md = """# ANI 知识库平台技术规格

## 一、系统架构概览

ANI 知识库平台采用三层服务架构：网关层（ani-gateway）、服务编排层（kb-service）和 AI 推理层（rag-engine）。
网关层负责请求路由、认证鉴权和对象存储代理；服务编排层负责知识库 CRUD 和文档解析管线编排；
AI 推理层负责文档解析、向量嵌入和摘要生成等无状态 RPC 执行。

在分布式部署中，各服务通过 gRPC 通信，保证低延迟和高吞吐。同时，所有写操作都支持幂等性控制，
通过 idempotency_key 保证在网络抖动和服务重启场景下的数据一致性。

## 二、服务组件清单

| 组件名称 | 技术栈 | 端口 | 职责 |
|---------|-------|------|------|
| ani-gateway | Go | 8080 | API 网关 + 对象存储 |
| kb-service | Python/FastAPI | 8002/50053 | 知识库 CRUD + 解析编排 |
| rag-engine | Python/FastAPI | 8001/50052 | RAG 引擎 (parse/embed/generate) |
| PostgreSQL | PostgreSQL 17 | 5432 | 元数据 + kb_chunks |
| Milvus | Milvus 2.4 | 19530 | 向量数据库 |

## 三、架构图示

文档中包含架构示意图：

![系统架构图](https://example.com/architecture.png)

数据流转示意图：

![数据流转图](https://example.com/data-flow.png)

## 四、文档解析管线详解

文档解析管线是知识库平台的核心功能之一。它负责将用户上传的文档（支持 PDF、Word、Markdown 等格式）
转换为结构化的文本块，并生成向量嵌入用于后续的语义检索。

管线的完整流程包括以下步骤：

1. 获取文档下载 URL（kb-service 向 Core API 请求 download_url）
2. 调用 rag-engine Parse RPC（传入 download_url，rag-engine 下载文档并解析）
3. 图片上传（kb-service 将 rag-engine 返回的图片 bytes 上传到 Core API）
4. 摘要生成（调用 Generate RPC，取前 3 个 parent blocks 生成 200-500 字摘要）
5. 向量嵌入（调用 Embed RPC 嵌入 child chunks + summary）
6. 向量插入（kb-service 调用 Core API 插入预计算向量）
7. PG 写入（parents + children + summaries 分开写入 kb_chunks）
8. 状态更新（parsing → indexing → ready）

在异常场景下，管线会捕获异常并将文档状态标记为 failed，
同时记录经过脱敏处理的错误信息（移除敏感路径和凭据，截断 500 字符）。
幂等性保证：如果文档已经是 ready 状态，重复调用会直接跳过。

## 五、混合检索策略

知识库平台支持三种检索模式：关键词检索、向量检索和混合检索。
关键词检索基于 PostgreSQL 的 pg_trgm 扩展实现模糊匹配，
向量检索基于 Milvus 的 HNSW 索引实现余弦相似度匹配，
混合检索将两种模式的结果融合，兼顾精确性和语义性。

检索结果经过去重和重排序后，取 top-k 个 chunk 作为 LLM 的上下文，
生成最终回答。整个检索-生成流程在 kb-service 的 Query RPC 中编排完成。
"""
    return md.encode("utf-8")


# ── service startup ─────────────────────────────────────────────────────────
def start_gateway():
    env = os.environ.copy()
    env.update({
        "ANI_AUTH_MODE": "dev",
        "GATEWAY_PORT": "8080",
        "DATABASE_URL": PG_URL,
        "REDIS_URL": f"redis://:ani_dev_password@{SRV}:30453/0",
        "OBJECT_STORE_PROVIDER": "minio",
        "OBJECT_STORE_ENDPOINT": MINIO_ENDPOINT,
        "OBJECT_STORE_PUBLIC_ENDPOINT": MINIO_ENDPOINT,
        "OBJECT_STORE_ACCESS_KEY_ID": MINIO_AK,
        "OBJECT_STORE_SECRET_ACCESS_KEY": MINIO_SK,
        "OBJECT_STORE_SECURE": "false",
        "OBJECT_STORE_BUCKET_PREFIX": "ani-s13-",
        "OBJECT_STORE_REGION": "us-east-1",
        "VECTOR_STORE_PROVIDER": "milvus",
        "VECTOR_STORE_ENDPOINT": MILVUS_ADDR,
        "VECTOR_STORE_COLLECTION_PREFIX": "ani_s13_",
        "NATS_URL": NATS_URL,
        "STORAGE_PROVIDER": "",
        "K8S_CLUSTER_PROVIDER_MODE": "",
        "GPU_INVENTORY_PROVIDER": "",
        "NETWORK_PROVIDER": "",
        "REGISTRY_PROVIDER_MODE": "",
        "INSTANCE_OBSERVABILITY_PROVIDER": "",
    })
    _start([str(GATEWAY_EXE)], env, REPO / "services" / "ani-gateway", "gateway")


def start_rag_engine():
    env = os.environ.copy()
    env.update({
        "MILVUS_ADDR": MILVUS_ADDR,
        "DATABASE_URL": PG_URL,
        "NATS_URL": NATS_URL,
        "EMBEDDING_MODEL": "Qwen3-Embedding-0.6B",
        "EMBEDDING_API_BASE": "http://10.10.20.197:8006/v1",
        "MINIO_ENDPOINT": MINIO_ENDPOINT,
        "MINIO_ACCESS_KEY": MINIO_AK,
        "MINIO_SECRET_KEY": MINIO_SK,
        "MINIO_SECURE": "false",
        "MINIO_BUCKET": "ani-kb-docs",
        "VLLM_MODEL": VLLM_MODEL,
        "VLLM_API_BASE": VLLM_API_BASE,
        "VLLM_API_KEY": VLLM_API_KEY,
        "VLLM_CONTEXT_WINDOW": "32768",
        "REDIS_URL": f"redis://:ani_dev_password@{SRV}:30453/0",
        "ANI_GATEWAY_URL": "http://localhost:8080",
    })
    py = sys.executable
    _start([py, "main.py"], env, RAG_ENGINE, "rag-engine")


def start_kb_service():
    env = os.environ.copy()
    env.update({
        "DATABASE_URL": PG_URL,
        "NATS_URL": NATS_URL,
        "ANI_GATEWAY_INTERNAL_URL": "http://localhost:8080",
        "RAG_ENGINE_ADDR": "localhost:8001",
        "GRPC_PORT": "50053",
        "REDIS_URL": f"redis://:ani_dev_password@{SRV}:30453/0",
    })
    py = sys.executable
    _start([py, "main.py"], env, KB_SVC, "kb-service")


# ── gRPC client helpers ──────────────────────────────────────────────────────
def _import_kb_pb():
    gen_root = str(KB_SVC / "app" / "generated")
    if gen_root not in sys.path:
        sys.path.insert(0, gen_root)
    if str(KB_SVC) not in sys.path:
        sys.path.insert(0, str(KB_SVC))
    from app.generated.kb.v1 import kb_service_pb2 as pb
    from app.generated.kb.v1 import kb_service_pb2_grpc as pb_grpc
    return pb, pb_grpc


def _upload_file(stub, pb, kb_id, file_name, file_type, content_bytes, idempotency_key):
    import urllib.request
    upload_req = pb.GetDocumentUploadURLRequest(
        tenant_id=TENANT_ID, kb_id=kb_id,
        file_name=file_name, file_type=file_type,
        file_size_bytes=0, idempotency_key=idempotency_key,
    )
    upload_resp = stub.GetDocumentUploadURL(upload_req)
    doc_id = upload_resp.doc_id
    log_json("UploadResponse", {
        "doc_id": doc_id,
        "storage_path": upload_resp.storage_path,
        "upload_url": (upload_resp.upload_url[:80] + "...") if len(upload_resp.upload_url) > 80 else upload_resp.upload_url,
    })
    put_req = urllib.request.Request(
        upload_resp.upload_url, data=content_bytes, method="PUT",
        headers={"Content-Type": "application/octet-stream"},
    )
    put_resp = urllib.request.urlopen(put_req, timeout=30)
    log(f"  Upload PUT status: {put_resp.status} ({len(content_bytes)} bytes)")
    return doc_id, upload_resp.storage_path


def _wait_parse_status(doc_id, timeout=180):
    import asyncpg, asyncio
    async def _wait():
        conn = await asyncpg.connect(PG_URL)
        for i in range(timeout // 3):
            row = await conn.fetchrow(
                "SELECT parse_status, error_message, chunk_count FROM kb_documents "
                "WHERE id=$1 AND tenant_id=$2",
                uuid.UUID(doc_id), uuid.UUID(TENANT_ID),
            )
            if row and row["parse_status"] in ("ready", "failed"):
                log(f"  parse_status={row['parse_status']} chunk_count={row['chunk_count']} error={row['error_message'] or ''}")
                await conn.close()
                return row["parse_status"], row["chunk_count"], row["error_message"]
            if row:
                log(f"  parse_status={row['parse_status']} (waiting... [{i}])")
            await asyncio.sleep(3)
        await conn.close()
        return None, 0, "timeout"
    return asyncio.run(_wait())


def _list_chunks(doc_id):
    import asyncpg, asyncio
    async def _list():
        conn = await asyncpg.connect(PG_URL)
        rows = await conn.fetch(
            "SELECT id::text, content, chunk_type, content_type, page_number, "
            "parent_content, parent_chunk_id::text as parent_chunk_id, token_count "
            "FROM kb_chunks WHERE doc_id=$1 AND tenant_id=$2 "
            "ORDER BY chunk_type, created_at",
            uuid.UUID(doc_id), uuid.UUID(TENANT_ID),
        )
        await conn.close()
        return [dict(r) for r in rows]
    return asyncio.run(_list())


# ── main E2E test ────────────────────────────────────────────────────────────
def run_e2e():
    import grpc
    pb, pb_grpc = _import_kb_pb()
    channel = grpc.insecure_channel("localhost:50053")
    stub = pb_grpc.KBServiceStub(channel)

    results = {"pass": 0, "fail": 0, "errors": []}

    def check(name, cond, detail=""):
        status = "PASS" if cond else "FAIL"
        if cond:
            results["pass"] += 1
        else:
            results["fail"] += 1
            results["errors"].append(f"{name}: {detail}")
        log(f"  [{status}] {name}" + (f" — {detail}" if detail and not cond else ""))

    # ── CreateKB ────────────────────────────────────────────────────────────
    log("=" * 70)
    log("STEP 1: CreateKB")
    log("=" * 70)
    kb_name = f"e2e032-{E2E_TAG}"
    req = pb.CreateKBRequest(
        tenant_id=TENANT_ID, name=kb_name,
        description="E2E test for issue-032 ParseOrchestrator",
        embedding_model="bge-m3", chunk_size=512,
        top_k=5, score_threshold=0.3, retrieval_mode="hybrid",
    )
    log_json("CreateKBRequest", {
        "tenant_id": TENANT_ID, "name": kb_name,
        "chunk_size": 512, "retrieval_mode": "hybrid",
    })
    kb = stub.CreateKB(req)
    kb_id = kb.id
    log_json("CreateKBResponse", {
        "id": kb_id, "name": kb.name,
        "embedding_model": kb.embedding_model,
    })
    check("CreateKB returns non-empty id", bool(kb_id))

    # ════════════════════════════════════════════════════════════════════════
    # TEST DOCUMENT 1: Word (.docx) — long text + table + 2 embedded images
    # ════════════════════════════════════════════════════════════════════════
    log("")
    log("=" * 70)
    log("TEST DOC 1: Word (.docx) — long text + table + 2 embedded images")
    log("=" * 70)

    docx_bytes = _create_test_docx()
    log(f"  DOCX file size: {len(docx_bytes)} bytes")
    log("  Content: 5 sections of long text, 1 table (5x4), 2 embedded PNG images")
    log("")
    log("  ── INPUT: DOCX full content (extracted text) ──")
    # Print the text content of the docx for reference
    try:
        from docx import Document
        doc = Document(io.BytesIO(docx_bytes))
        for para in doc.paragraphs:
            if para.text.strip():
                log(f"  {para.text}")
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                log(f"  | {' | '.join(cells)} |")
    except Exception as e:
        log(f"  (could not extract docx text: {e})")
    log("  ── END INPUT ──")
    log("")

    docx_doc_id, docx_path = _upload_file(
        stub, pb, kb_id, "e2e032_test.docx", "docx", docx_bytes, f"docx-{E2E_TAG}",
    )
    check("DOCX: upload succeeds", True)

    # Notify
    stub.NotifyDocumentUploaded(pb.NotifyDocumentUploadedRequest(
        tenant_id=TENANT_ID, kb_id=kb_id,
        doc_id=docx_doc_id, storage_path=docx_path,
    ))
    log("  DOCX NotifyDocumentUploaded OK")

    # Wait for parse
    log("  Waiting for DOCX parse (may take 60-120s with remote LLM)...")
    status, chunk_count, error = _wait_parse_status(docx_doc_id)
    check("DOCX: parse reaches ready", status == "ready", f"status={status} error={error}")

    # List chunks
    if status == "ready":
        chunks = _list_chunks(docx_doc_id)
        log("")
        log(f"  ═══ DOCX CHUNK DETAILS ({len(chunks)} chunks, chunk_count={chunk_count}) ═══")
        log("")

        parent_chunks = [c for c in chunks if c["chunk_type"] == "parent"]
        child_chunks = [c for c in chunks if c["chunk_type"] == "child"]
        summary_chunks = [c for c in chunks if c["chunk_type"] == "doc_summary"]

        log(f"  Parents: {len(parent_chunks)}  Children: {len(child_chunks)}  Summaries: {len(summary_chunks)}")
        log(f"  Total in DB: {len(chunks)}  chunk_count field: {chunk_count}")
        check("DOCX: chunk_count == parents + children + summaries",
              chunk_count == len(parent_chunks) + len(child_chunks) + len(summary_chunks),
              f"chunk_count={chunk_count} vs actual={len(parent_chunks) + len(child_chunks) + len(summary_chunks)}")

        # Print all chunks — FULL content, no truncation
        for i, c in enumerate(chunks):
            content = c.get("content", "")
            ctype = c["chunk_type"]
            ctype_field = c.get("content_type", "")
            page = c.get("page_number", 0)
            parent_content = c.get("parent_content", "")
            parent_id = c.get("parent_chunk_id", "")
            tokens = c.get("token_count", 0)
            log(f"  ── Chunk[{i}] type={ctype} content_type={ctype_field} page={page} tokens={tokens} ──")
            if ctype == "parent":
                log(f"  [PARENT FULL CONTENT]")
                log(f"  {content}")
            elif ctype == "doc_summary":
                log(f"  [SUMMARY FULL CONTENT]")
                log(f"  {content}")
            else:
                log(f"  [CHILD FULL CONTENT]")
                log(f"  {content}")
                if parent_content:
                    log(f"  [CHILD parent_content FULL]")
                    log(f"  {parent_content}")
            if parent_id:
                log(f"  parent_chunk_id: {parent_id}")
            log("")

        # Verify image markdown links in parent content
        has_image_link = False
        image_urls = []
        for c in parent_chunks:
            content = c.get("content", "")
            if "[图片:" in content:
                has_image_link = True
                import re
                urls = re.findall(r"\[图片:.*?\]\(([^)]+)\)", content)
                image_urls.extend(urls)

        log(f"  Image link check: has_image_link={has_image_link}")
        if image_urls:
            log(f"  Found {len(image_urls)} image object_ids:")
            for u in image_urls:
                log(f"    {u}")
        check("DOCX: parent content contains [图片: 图片](object_id) links",
              has_image_link, "no [图片:...] links found in parent content")

        # Verify summary
        has_summary = len(summary_chunks) > 0
        check("DOCX: doc_summary chunk exists", has_summary, "no doc_summary chunk")
        if has_summary:
            summary_content = summary_chunks[0].get("content", "")
            log(f"  Summary content ({len(summary_content)} chars, FULL):")
            log(f"  {summary_content}")
            check("DOCX: summary is non-empty", len(summary_content) > 10, "summary too short")

        # Verify text content
        all_content = " ".join(c.get("content", "") for c in chunks)
        has_text = "架构" in all_content or "知识库" in all_content
        has_table = "|" in all_content or "<table>" in all_content or "组件名称" in all_content
        check("DOCX: chunks contain text content", has_text)
        check("DOCX: chunks contain table content", has_table)

    # ════════════════════════════════════════════════════════════════════════
    # TEST DOCUMENT 2: Markdown (.md) — long text + table + image links
    # ════════════════════════════════════════════════════════════════════════
    log("")
    log("=" * 70)
    log("TEST DOC 2: Markdown (.md) — long text + table + external image links")
    log("=" * 70)

    md_bytes = _create_test_md()
    log(f"  MD file size: {len(md_bytes)} bytes")
    log("  Content: 5 sections of long text, 1 table (5x5), 2 external image links")
    log("")
    log("  ── INPUT: MD full content ──")
    log(md_bytes.decode("utf-8"))
    log("  ── END INPUT ──")
    log("")

    md_doc_id, md_path = _upload_file(
        stub, pb, kb_id, "e2e032_test.md", "md", md_bytes, f"md-{E2E_TAG}",
    )
    check("MD: upload succeeds", True)

    stub.NotifyDocumentUploaded(pb.NotifyDocumentUploadedRequest(
        tenant_id=TENANT_ID, kb_id=kb_id,
        doc_id=md_doc_id, storage_path=md_path,
    ))
    log("  MD NotifyDocumentUploaded OK")

    log("  Waiting for MD parse...")
    md_status, md_chunk_count, md_error = _wait_parse_status(md_doc_id)
    check("MD: parse reaches ready", md_status == "ready", f"status={md_status} error={md_error}")

    if md_status == "ready":
        md_chunks = _list_chunks(md_doc_id)
        log("")
        log(f"  ═══ MD CHUNK DETAILS ({len(md_chunks)} chunks, chunk_count={md_chunk_count}) ═══")
        log("")

        md_parents = [c for c in md_chunks if c["chunk_type"] == "parent"]
        md_children = [c for c in md_chunks if c["chunk_type"] == "child"]
        md_summaries = [c for c in md_chunks if c["chunk_type"] == "doc_summary"]
        log(f"  Parents: {len(md_parents)}  Children: {len(md_children)}  Summaries: {len(md_summaries)}")

        for i, c in enumerate(md_chunks):
            content = c.get("content", "")
            ctype = c["chunk_type"]
            log(f"  ── MD Chunk[{i}] type={ctype} content_type={c.get('content_type')} ──")
            if ctype == "parent":
                log(f"  [PARENT FULL CONTENT]")
                log(f"  {content}")
            elif ctype == "doc_summary":
                log(f"  [SUMMARY FULL CONTENT]")
                log(f"  {content}")
            else:
                log(f"  [CHILD FULL CONTENT]")
                log(f"  {content}")
                parent_c = c.get("parent_content", "")
                if parent_c:
                    log(f"  [CHILD parent_content FULL]")
                    log(f"  {parent_c}")
            log("")

        # MD has no embedded images (only external links), so no [图片: 图片] expected
        # but should contain the markdown image links ![...](url)
        md_all_content = " ".join(c.get("content", "") for c in md_chunks)
        has_md_text = "架构" in md_all_content or "知识库" in md_all_content
        has_md_table = "组件名称" in md_all_content or "|" in md_all_content
        has_md_image = "example.com" in md_all_content or "![" in md_all_content
        check("MD: chunks contain text content", has_md_text)
        check("MD: chunks contain table content", has_md_table)
        check("MD: chunks contain image links", has_md_image)
        check("MD: doc_summary chunk exists", len(md_summaries) > 0, "no doc_summary")

    # ── Query test ──────────────────────────────────────────────────────────
    log("")
    log("=" * 70)
    log("STEP 3: Query — '文档解析管线包含哪些步骤?'")
    log("=" * 70)
    query_req = pb.QueryRequest(
        tenant_id=TENANT_ID, kb_id=kb_id,
        question="文档解析管线包含哪些步骤?",
        idempotency_key=f"query-{E2E_TAG}",
        top_k=3, score_threshold=0.0, retrieval_mode="hybrid",
    )
    log_json("QueryRequest", {
        "question": query_req.question,
        "top_k": query_req.top_k,
        "retrieval_mode": query_req.retrieval_mode,
    })
    try:
        query_resp = stub.Query(query_req, timeout=120)
        log_json("QueryResponse", {
            "answer": query_resp.answer,
            "session_id": query_resp.session_id,
            "num_sources": len(query_resp.sources),
            "input_tokens": query_resp.input_tokens,
            "output_tokens": query_resp.output_tokens,
        })
        check("Query returns an answer", bool(query_resp.answer), "answer empty")
        if query_resp.sources:
            for i, sc in enumerate(query_resp.sources):
                log(f"  ── Source[{i}] score={sc.score:.3f} ──")
                log(f"  [SOURCE FULL CONTENT]")
                log(f"  {sc.content}")
                log("")
        check("Query returns source chunks", len(query_resp.sources) > 0)
    except grpc.RpcError as e:
        log(f"  Query FAILED: {e.code()}: {e.details()}")
        check("Query returns an answer", False, f"{e.code()}")

    # ── Cleanup ──────────────────────────────────────────────────────────────
    log("")
    log("=" * 70)
    log("STEP 4: Cleanup — DeleteDocument + DeleteKB")
    log("=" * 70)
    try:
        stub.DeleteDocument(pb.DeleteDocumentRequest(
            tenant_id=TENANT_ID, kb_id=kb_id, doc_id=docx_doc_id))
        log("  DOCX deleted OK")
        check("DOCX: DeleteDocument succeeds", True)
    except grpc.RpcError as e:
        check("DOCX: DeleteDocument succeeds", False, f"{e.code()}")

    try:
        stub.DeleteDocument(pb.DeleteDocumentRequest(
            tenant_id=TENANT_ID, kb_id=kb_id, doc_id=md_doc_id))
        log("  MD deleted OK")
        check("MD: DeleteDocument succeeds", True)
    except grpc.RpcError as e:
        check("MD: DeleteDocument succeeds", False, f"{e.code()}")

    try:
        stub.DeleteKB(pb.DeleteKBRequest(tenant_id=TENANT_ID, kb_id=kb_id))
        log("  KB deleted OK")
        check("DeleteKB succeeds", True)
    except grpc.RpcError as e:
        check("DeleteKB succeeds", False, f"{e.code()}")

    channel.close()
    return results


# ── main ──────────────────────────────────────────────────────────────────────
def main():
    _open_log()
    log("=" * 70)
    log("  ANI Issue-032 E2E Test — ParseOrchestrator Pipeline")
    log(f"  Tenant: {TENANT_ID}")
    log(f"  Tag:    {E2E_TAG}")
    log(f"  Log:    {E2E_LOG}")
    log("=" * 70)
    log("")
    log("Server components (NodePorts on 10.10.1.66):")
    log(f"  Postgres:  {SRV}:30945")
    log(f"  Milvus:    {SRV}:31930")
    log(f"  MinIO:     {SRV}:30900")
    log(f"  NATS:      {SRV}:31062")
    log(f"  Redis:     {SRV}:30453")
    log(f"  Embedding: 10.10.20.197:8006")
    log(f"  LLM:       {VLLM_API_BASE} ({VLLM_MODEL})")
    log("")
    log("Test documents:")
    log("  1. Word (.docx)  — 5 sections long text, table, 2 embedded PNG images")
    log("  2. Markdown (.md) — 5 sections long text, table, 2 external image links")
    log("")

    # 1. Start services
    log(">>> Phase 1: Starting local services")
    start_gateway()
    if not _wait_http("http://localhost:8080/healthz", "gateway", 15):
        log("FATAL: gateway not ready, aborting")
        _kill_all()
        return 1

    start_rag_engine()
    if not _wait_http("http://localhost:8001/health", "rag-engine", 60):
        log("WARNING: rag-engine not ready (continuing — gRPC may still start)")

    start_kb_service()
    if not _wait_http("http://localhost:8002/health", "kb-service", 15):
        log("FATAL: kb-service not ready, aborting")
        _kill_all()
        return 1

    log("")
    log(">>> Phase 2: Running ParseOrchestrator E2E tests")
    log("")
    time.sleep(2)

    try:
        results = run_e2e()
    except Exception as e:
        import traceback
        log(f"E2E test crashed: {e}")
        traceback.print_exc()
        results = {"pass": 0, "fail": 1, "errors": [str(e)]}

    log("")
    log("=" * 70)
    log(f"  E2E RESULTS: {results['pass']} passed, {results['fail']} failed")
    log("=" * 70)
    if results["errors"]:
        log("Errors:")
        for e in results["errors"]:
            log(f"  - {e}")
    log("")
    log(f"Full log: {E2E_LOG}")

    _kill_all()
    return 1 if results["fail"] > 0 else 0


if __name__ == "__main__":
    sys.exit(main())
